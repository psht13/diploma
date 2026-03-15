from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

import fitz
from pypdf import PdfReader

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "assets" / "coursework_draft_ua_submission_ready.docx"
TARGET_DOC = ROOT / "output" / "coursework_draft_ua_submission_ready_v2.docx"
WORK_DIR = ROOT / "tmp" / "docs" / "submission_ready_v2"
STAGE_DOC = WORK_DIR / "coursework_submission_stage_v2.docx"
PDF_PATH = WORK_DIR / "coursework_draft_ua_submission_ready_v2.pdf"
PNG_DIR = WORK_DIR / "png"
DIAGRAM_DIR = ROOT / "docs" / "codex" / "diagram_assets"

UK_LANG = "uk-UA"
EN_LANG = "en-US"
CUSTOM_STYLES = {
    "ListingCaption",
    "SourceNote",
    "PromptBlock",
    "CodeBlock",
    "DialogueBlock",
    "AppendixFigureLabel",
}
TOC_LINES = [
    "АНОТАЦІЯ",
    "ABSTRACT",
    "СПИСОК УМОВНИХ ПОЗНАЧЕНЬ",
    "ВСТУП",
    "РОЗДІЛ 1. АНАЛІЗ ПРЕДМЕТНОЇ ГАЛУЗІ АДАПТИВНИХ ІНТЕЛЕКТУАЛЬНИХ НАВЧАЛЬНИХ СИСТЕМ",
    "1.1. Поняття та класифікація інтелектуальних навчальних систем (ITS/АІНС)",
    "1.2. Архітектурні компоненти АІНС",
    "1.3. Діалогові агенти у навчальних системах",
    "1.4. Огляд сучасних систем-репетиторів на основі ШІ",
    "1.5. Висновки до розділу 1",
    "РОЗДІЛ 2. ВЕЛИКІ МОВНІ МОДЕЛІ ЯК ОСНОВА ДІАЛОГОВОГО АГЕНТА",
    "2.1. Архітектура трансформерів та принципи роботи LLM",
    "2.2. Порівняння LLM для навчальних застосунків",
    "2.3. Проєктування промптів, RAG і донавчання як методи адаптації",
    "2.4. Обмеження LLM: галюцинації, пам'ять, етика, приватність",
    "2.5. Обґрунтування вибору локальної відкритої Llama для прототипу",
    "РОЗДІЛ 3. ПРОЄКТУВАННЯ АРХІТЕКТУРИ АІНС",
    "3.1. Функціональні та нефункціональні вимоги",
    "3.2. Загальна архітектура системи",
    "3.3. Модель студента",
    "3.4. Модель знань і навчального контенту",
    "3.5. Модель і стратегії репетитора",
    "3.5.1. Формальне визначення агента-репетитора",
    "3.5.2. Архітектура агента - компонентна модель",
    "3.5.3. Матриця рішень і стратегія мінімальної підказки",
    "3.5.4. Типові сценарії взаємодії",
    "3.5.5. Цикл роботи агента: міркування → дія → спостереження",
    "3.5.6. Організація пам'яті агента",
    "3.5.7. Роль LLM і її обмеження",
    "РОЗДІЛ 4. РЕАЛІЗАЦІЯ ТА ЕКСПЕРИМЕНТАЛЬНА ПЕРЕВІРКА",
    "4.1. Вибір технологічного стеку",
    "4.2. Реалізація ключових модулів системи",
    "4.3. Проєктування промптів для генерації вправ і фідбеку",
    "4.4. Дизайн та проведення експерименту",
    "4.4.1. Функціональна валідація прототипу",
    "4.4.2. Дизайн педагогічного експерименту",
    "4.5. Аналіз результатів",
    "ВИСНОВКИ",
    "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ",
    "ДОДАТКИ",
    "Додаток А. Діаграми архітектури системи",
    "Додаток Б. Приклади промптів",
    "Додаток В. Приклади діалогів агента",
    "Додаток Г. Фрагменти програмного коду",
]


def set_run_font(run, size: float = 14, bold: bool | None = None, italic: bool | None = None, font_name: str = "Times New Roman") -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def set_language(element, lang: str) -> None:
    lang_el = element.find(qn("w:lang"))
    if lang_el is None:
        lang_el = OxmlElement("w:lang")
        element.append(lang_el)
    lang_el.set(qn("w:val"), lang)
    lang_el.set(qn("w:eastAsia"), lang)
    lang_el.set(qn("w:bidi"), lang)


def set_paragraph_language(paragraph: Paragraph, lang: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_r_pr = p_pr.find(qn("w:rPr"))
    if p_r_pr is None:
        p_r_pr = OxmlElement("w:rPr")
        p_pr.append(p_r_pr)
    set_language(p_r_pr, lang)
    for run in paragraph.runs:
        set_language(run._element.get_or_add_rPr(), lang)


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("pPr"):
            continue
        paragraph._p.remove(child)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    style_name: str | None = None,
    size: float = 14,
    bold: bool | None = None,
    italic: bool | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    font_name: str = "Times New Roman",
    lang: str = UK_LANG,
) -> None:
    clear_paragraph(paragraph)
    if style_name is not None:
        paragraph.style = style_name
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, font_name=font_name)
    if align is not None:
        paragraph.alignment = align
    set_paragraph_language(paragraph, lang)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = style_name
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_content_between(start_paragraph: Paragraph, end_paragraph: Paragraph) -> None:
    node = start_paragraph._p.getnext()
    while node is not None and node != end_paragraph._p:
        next_node = node.getnext()
        remove_element(node)
        node = next_node


def find_paragraph_exact(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_paragraph_by_prefix(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def get_style(document: Document, name: str):
    for style in document.styles:
        if style.name == name:
            return style
    raise KeyError(f"Style not found: {name}")


def add_or_update_style(document: Document, name: str, base_style: str = "Normal"):
    styles = document.styles
    try:
        style = get_style(document, name)
    except KeyError:
        style = None

    if style is not None:
        pass
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = get_style(document, base_style)
    return style


def shade_paragraph(paragraph: Paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def format_title_topic(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Mm(0)
    label = paragraph.add_run("Тема: ")
    set_run_font(label, size=16, bold=True)
    topic = paragraph.add_run("“Розробка архітектури АІНС на основі великих мовних моделей для реалізації діалогового агента з функціями репетитора”")
    set_run_font(topic, size=16, bold=False)
    topic.font.underline = True
    set_paragraph_language(paragraph, UK_LANG)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(215.9)
    section.page_height = Mm(279.4)
    section.left_margin = Mm(20)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    normal = get_style(document, "Normal")
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Mm(12.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    set_language(normal._element.get_or_add_rPr(), UK_LANG)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number", "Title"]:
        style = get_style(document, style_name)
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = style_name.startswith("Heading") or style_name == "Title"
        set_language(style._element.get_or_add_rPr(), UK_LANG)

    style = add_or_update_style(document, "ListingCaption")
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.bold = True
    style.paragraph_format.first_line_indent = Mm(0)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(3)

    style = add_or_update_style(document, "SourceNote")
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.font.italic = True
    style.paragraph_format.first_line_indent = Mm(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)

    style = add_or_update_style(document, "PromptBlock")
    style.font.name = "Courier New"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.first_line_indent = Mm(0)
    style.paragraph_format.left_indent = Mm(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)

    style = add_or_update_style(document, "CodeBlock")
    style.font.name = "Courier New"
    style.font.size = Pt(9.5)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.first_line_indent = Mm(0)
    style.paragraph_format.left_indent = Mm(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)

    style = add_or_update_style(document, "DialogueBlock")
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.first_line_indent = Mm(-10)
    style.paragraph_format.left_indent = Mm(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.0

    style = add_or_update_style(document, "AppendixFigureLabel")
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.bold = True
    style.paragraph_format.first_line_indent = Mm(0)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(3)

    for style_name in CUSTOM_STYLES:
        set_language(get_style(document, style_name)._element.get_or_add_rPr(), UK_LANG)


def prepare_toc_region(document: Document) -> None:
    toc_heading = find_paragraph_exact(document, "ЗМІСТ")
    next_heading = find_paragraph_exact(document, "СПИСОК УМОВНИХ ПОЗНАЧЕНЬ")

    toc_paragraphs: list[Paragraph] = []
    node = toc_heading._p.getnext()
    while node is not None and node != next_heading._p:
        if node.tag.endswith("}p"):
            toc_paragraphs.append(Paragraph(node, document))
        node = node.getnext()

    for index, toc_line in enumerate(TOC_LINES):
        paragraph = toc_paragraphs[index]
        paragraph.style = "Normal"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Mm(0)
        paragraph.paragraph_format.left_indent = Mm(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        set_paragraph_text(paragraph, f"{toc_line}\t0", lang=EN_LANG if toc_line == "ABSTRACT" else UK_LANG)

    for paragraph in toc_paragraphs[len(TOC_LINES):]:
        remove_element(paragraph._element)


def replace_text_content(document: Document) -> None:
    title_paragraph = find_paragraph_by_prefix(document, "Тема:")
    format_title_topic(title_paragraph)

    replacements = {
        "The document separates conceptual architecture from practical implementation.":
            "The document separates conceptual architecture from practical implementation. "
            "A local Ollama-based execution environment, Web Worker sandboxing, JSON retry and "
            "backup logic, and an honest experimental design without fabricated educational "
            "outcomes are treated as core principles.",
        "У центрі персоналізації перебуває модель студента.":
            "У центрі персоналізації перебуває модель студента. Повноцінні ITS використовують "
            "для цього складніші ймовірнісні моделі або механізми трасування знань, але для "
            "мінімального практичного сценарію доцільно застосувати спрощений вектор стану "
            "UserState, який усе ж зберігає ключові сигнали для педагогічного рішення [2], [3], [4].",
        "У прототипі ця декомпозиція реалізована не повністю:":
            "У прототипі ця декомпозиція реалізована не повністю: роль модуля сприйняття "
            "виконують явний режим пояснення, представлений параметром interactionMode зі "
            "значенням explain, та прості текстові евристики для вільного запиту, а роль модуля "
            "генерації виконує локальна LLM або резервний механізм. Саме така схема "
            "використовується як основа повнішої системи в магістерській проєкції.",
        "Ключова вимога до агента-репетитора полягає":
            "Ключова вимога до агента-репетитора полягає в тому, щоб допомога була мінімально "
            "достатньою. Після технічного аудиту правило ескалації було уточнено: перша невдала "
            "спроба приводить до мінімальної підказки (minimal_hint), а адресна підказка "
            "(targeted_hint) з'являється лише після повторення тієї самої помилки або серії "
            "кількох невдалих спроб.",
        "У коді ця матриця реалізована в модулі":
            "У коді ця матриця реалізована в модулі TutorPolicy через набір явних правил. Для "
            "більш складної системи вона може бути замінена рушієм правил або адаптивною "
            "стратегією вибору дії, але для курсової роботи важливо насамперед показати сам "
            "принцип явного розділення педагогічної дії та генеративної реалізації.",
        "Повторна помилка. Агент переходить":
            "Повторна помилка. Агент переходить до більш адресної підказки "
            "(targeted_hint) і конкретніше називає, який фрагмент логіки слід перевірити.",
        "Запит на пояснення. Агент переходить":
            "Запит на пояснення. Агент переходить у режим концептуального пояснення "
            "(concept_explanation) і дає теоретичне пояснення без завершеного розв'язку.",
        "Робочий цикл агента описується схемою":
            "Робочий цикл агента описується схемою «міркування → дія → спостереження», яка "
            "відповідає підходу ReAct. На етапі міркування система інтерпретує нові "
            "спостереження й оновлює UserState. На етапі дії обирається педагогічна дія: "
            "створити вправу, дати підказку, пояснити концепт або запросити нову спробу. На "
            "етапі спостереження система отримує нові результати, зокрема результати тестів, і "
            "повторює цикл [12].",
        "На рисунку 4.1 показано, як у прототипі":
            "На рисунку 4.1 показано, як у прототипі розведено браузерний інтерфейс, локальний "
            "LLM-виклик, резервний режим і клієнтський запуск тестів у Web Worker.",
        "У поточній роботі виконано лише ту функціональну перевірку":
            "У поточній роботі виконано лише ту функціональну перевірку, яку можна підтвердити "
            "реальними артефактами середовища: модульні тести (npm test), контрольну перевірку "
            "(npm run check), браузерний прогін через Playwright і реальний локальний виклик "
            "Ollama + llama3.1:8b. Освітні результати, UX-оцінки та ефективність навчання тут не "
            "вигадуються. Натомість підтверджено коректну ескалацію підказок, явний режим "
            "пояснення, генерацію вправи локальною моделлю, полегшений JSON-експорт сесії та "
            "стійкість прототипу до невалідних LLM-відповідей.",
        "Отже, інженерна функціональність мінімального практичного сценарію підтверджена":
            "Отже, інженерна функціональність мінімального практичного сценарію підтверджена як "
            "модульними тестами, так і повним браузерним сценарієм із локальною Ollama. Під час "
            "цього прогону було виявлено й усунуто два локальні недоліки: збільшено тайм-аут "
            "LLM-виклику для llama3.1:8b і виправлено скидання firstFailure під час переходу до "
            "нової вправи.",
        "На момент завершення роботи наявні лише результати інженерної валідації:":
            "На момент завершення роботи наявні лише результати інженерної валідації: модульні "
            "тести, контрольний сценарій npm run check і браузерний прогін через Playwright з "
            "локальною Ollama + llama3.1:8b.",
        "Для практичної частини реалізовано мінімальний практичний сценарій":
            "Для практичної частини реалізовано мінімальний практичний сценарій репетиторського "
            "супроводу з JavaScript: генерація вправи, введення коду, запуск тестів у Web Worker, "
            "адаптивний фідбек, SessionMemory, явний режим пояснення та резервна логіка у "
            "випадку недоступності або нестабільності локальної LLM.",
    }

    for paragraph in document.paragraphs:
        for prefix, replacement in replacements.items():
            if paragraph.text.startswith(prefix):
                paragraph.text = replacement
                break

        if paragraph.style.name in CUSTOM_STYLES:
            continue

        if paragraph.text:
            paragraph.text = (
                paragraph.text.replace("`", "")
                .replace("LLM-based", "LLM-орієнтованих")
                .replace("retrieval-шаром", "шаром пошуку")
                .replace("AI tutor", "ШІ-репетитор")
                .replace("GPT-4-based UX", "інтерфейс на основі GPT-4")
                .replace("Roleplay, Explain My Answer, ", "рольова взаємодія, пояснення відповіді, ")
                .replace("retry/fallback", "повторні спроби та резервний режим")
                .replace("rule-based", "на основі правил")
                .replace("transcript tail", "кінцева частина журналу взаємодії")
                .replace("->", "→")
            )

    formula_explanation = find_paragraph_by_prefix(document, "Тут:")
    formula_explanation.text = (
        "Тут: K - інтегрована оцінка засвоєння теми; E - історія типових помилок; "
        "C - робоча оцінка впевненості; Topic - поточна тема; Attempts - кількість спроб; "
        "M - пам'ять сесії; O - нові спостереження; A - множина педагогічних дій; "
        "T - функція переходу стану; P - політика вибору дії; R - функція формування "
        "відповіді; Ctx - контекст генерації; Text - текстова відповідь агента. Такий запис "
        "відокремлює стан, політику та мовну генерацію і робить архітектуру агента придатною "
        "для формального опису без ототожнення її з однією лише LLM [11], [12]."
    )


def normalize_paragraphs(document: Document) -> None:
    annotation_index = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text == "АНОТАЦІЯ")

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name
        is_formula = text.startswith(("S = ", "A = {", "T: ", "P: ", "R: ")) or text.startswith("      ")

        if index < annotation_index:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                size = 16 if index in {0, 6, 8} else 14
                set_run_font(run, size=size, font_name="Times New Roman")
            set_paragraph_language(paragraph, UK_LANG)
            continue

        if style_name in CUSTOM_STYLES:
            if style_name in {"PromptBlock", "CodeBlock"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Mm(0)
                paragraph.paragraph_format.left_indent = Mm(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.keep_together = True
                shade_paragraph(paragraph, "F2F2F2")
                for run in paragraph.runs:
                    set_run_font(run, size=10.5 if style_name == "PromptBlock" else 9.5, font_name="Courier New")
            elif style_name == "DialogueBlock":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Mm(-10)
                paragraph.paragraph_format.left_indent = Mm(10)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    set_run_font(run, size=14, bold=bool(run.bold))
            elif style_name in {"ListingCaption", "AppendixFigureLabel"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Mm(0)
                paragraph.paragraph_format.keep_with_next = True
                for run in paragraph.runs:
                    set_run_font(run, size=12, bold=True)
            elif style_name == "SourceNote":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Mm(0)
                for run in paragraph.runs:
                    set_run_font(run, size=11, italic=True)
            continue

        if style_name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, size=14, bold=True)
        elif style_name in {"Heading 2", "Heading 3"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, size=14, bold=True)
        elif text == "ЗМІСТ" or "\t" in paragraph.text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if text else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.left_indent = Mm(0)
            for run in paragraph.runs:
                set_run_font(run, size=14)
        elif text.startswith("Таблиця "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, size=14)
        elif text.startswith("Рисунок "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            for run in paragraph.runs:
                set_run_font(run, size=14)
        elif is_formula:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, size=12)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Mm(12.5)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            for run in paragraph.runs:
                set_run_font(run, size=14, bold=bool(run.bold), italic=bool(run.italic))


def format_title_table(document: Document) -> None:
    title_table = document.tables[0]
    for row in title_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Mm(0)
                paragraph.paragraph_format.left_indent = Mm(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=12)
                set_paragraph_language(paragraph, UK_LANG)


def update_tables(document: Document) -> None:
    table = document.tables[1]
    table.cell(1, 2).text = "Діалоговий ШІ-репетитор і асистент викладача"
    table.cell(1, 3).text = "Підходить як приклад продуктового ШІ-репетитора [23]"
    table.cell(2, 2).text = "рольова взаємодія, пояснення відповіді, інтерфейс на основі GPT-4"
    table.cell(4, 3).text = "Показують сильні й слабкі сторони LLM у навчанні програмування [15], [17], [18], [26]"

    table = document.tables[2]
    table.cell(3, 1).text = "Запустити тести лише на клієнті, у Web Worker з тайм-аутом"
    table.cell(6, 1).text = "Обробити невалідний JSON від моделі через повторні спроби та резервний режим"

    table = document.tables[3]
    table.cell(1, 1).text = "Використовувати відкрите локальне середовище виконання"
    table.cell(1, 2).text = "цільове середовище виконання: Ollama + llama3.1:8b"
    table.cell(2, 1).text = "Не виконувати код студента на сервері"
    table.cell(4, 2).text = "виконано через модуль TutorPolicy"

    table = document.tables[4]
    table_rows = [
        ("Інтерфейс студента (Chat UI)", "Інтерфейс студента, показ вправи, коду, результатів тестів і фідбеку", "реалізовано"),
        ("Модуль сприйняття (Perception Module)", "Класифікація наміру й нормалізація запиту", "описано концептуально"),
        ("Сесійна пам'ять (SessionMemory)", "Поточний стан, кінцева частина журналу взаємодії, історія помилок", "реалізовано"),
        ("Пам'ять стислих підсумків (Summary Memory)", "Стисле зведення попередніх сесій", "описано концептуально"),
        ("Профільна пам'ять (Vector/Profile Memory)", "Довготривалий профіль студента між сесіями", "описано концептуально"),
        ("Пам'ять знань (Knowledge Memory)", "Контентна база, RAG, концепти, пояснення", "описано концептуально"),
        ("Модуль педагогічного вибору (TutorPolicy / Reasoning Module)", "Вибір педагогічної дії", "реалізовано у вигляді системи правил"),
        ("Генератор вправ (ExerciseGenerator)", "Створення вправи через LLM або резервну вправу", "реалізовано"),
        ("Запускач тестів (ClientTestRunner)", "Запуск тестів у Web Worker", "реалізовано"),
        ("Оцінювач фідбеку (FeedbackEvaluator)", "Формування адаптивного фідбеку", "реалізовано"),
        ("Модуль генерації (Generation Module / LLM)", "Генерація структурованих JSON-відповідей для вправ і фідбеку", "реалізовано частково через Ollama API"),
    ]
    for row_index, row_values in enumerate(table_rows, start=1):
        for column_index, value in enumerate(row_values):
            table.cell(row_index, column_index).text = value

    table = document.tables[5]
    table.cell(0, 0).text = "Поле моделі стану"
    state_rows = [
        "Рівень знань (knowledgeLevel)",
        "Історія помилок (errorHistory)",
        "Рівень упевненості (confidence)",
        "Поточна тема (currentTopic)",
        "Кількість спроб (attemptsCount)",
        "Остання дія агента (lastAction)",
    ]
    for row_index, value in enumerate(state_rows, start=1):
        table.cell(row_index, 0).text = value

    table = document.tables[6]
    actions = [
        "підкріплення з рефлексією (celebrate_and_reflect)",
        "мінімальна підказка (minimal_hint)",
        "адресна підказка (targeted_hint)",
        "усунення синтаксичної помилки (syntax_repair)",
        "концептуальне пояснення (concept_explanation)",
    ]
    for row_index, value in enumerate(actions, start=1):
        table.cell(row_index, 1).text = value

    table = document.tables[7]
    memory_rows = [
        "Буферна сесійна пам'ять (Buffer / Session Memory)",
        "Пам'ять стислих підсумків (Summary Memory)",
        "Профільна пам'ять (Vector / Profile Memory)",
        "Пам'ять знань (Knowledge Memory)",
    ]
    for row_index, value in enumerate(memory_rows, start=1):
        table.cell(row_index, 0).text = value
    table.cell(3, 1).text = "Стійкий профіль студента між сесіями, embeddings і патерни труднощів"

    table = document.tables[8]
    table.cell(3, 1).text = "SessionMemory у пам'яті процесу"
    table.cell(4, 1).text = "Web Worker + тайм-аут"
    table.cell(5, 1).text = "node --test"

    table = document.tables[9]
    rows = [
        ("Інтерфейс користувача", "prototype/index.html, prototype/src/main.js", "Інтерфейс вибору теми, редактор коду, журнал взаємодій, синхронізація моделі та дружні повідомлення про помилки."),
        ("Сесійна пам'ять (SessionMemory)", "prototype/src/state/sessionMemory.js", "Зберігає UserState, поточну вправу, історію повідомлень, першу зафіксовану помилку та формує полегшений JSON-експорт сесії."),
        ("Педагогічна політика (TutorPolicy)", "prototype/src/services/tutorPolicy.js", "Політика на основі правил з коректною ескалацією: перша помилка → minimal_hint, повторна та сама → targeted_hint, а кнопка пояснення через interactionMode = explain завжди веде до concept_explanation."),
        ("Генератор вправ (ExerciseGenerator)", "prototype/src/services/exerciseGenerator.js", "Генерація вправи через LLM, явна перевірка структури та перехід до резервної вправи при невідповідному JSON."),
        ("Запускач тестів (ClientTestRunner)", "prototype/src/services/clientTestRunner.js, prototype/src/workers/jsSandboxWorker.js", "Клієнтський запуск тестів у Web Worker з тайм-аутом, перевіркою вхідних даних, worker.onerror і messageerror."),
        ("Оцінювач фідбеку (FeedbackEvaluator)", "prototype/src/services/feedbackEvaluator.js", "Адаптивний фідбек через LLM або резервний режим; резервне пояснення формує змістовне пояснення концепту."),
        ("Клієнт LLM (OllamaClient)", "prototype/src/services/ollamaClient.js", "HTTP API, перевірка доступності, перевірка структури та повторні спроби з переходом до резервного режиму для JSON-відповідей."),
    ]
    for row_index, row_values in enumerate(rows, start=1):
        for column_index, value in enumerate(row_values):
            table.cell(row_index, column_index).text = value

    table = document.tables[10]
    checks = [
        ("Політика підказок", "модульні тести (npm test)", "пройдено", "Перевірено: перша помилка → minimal_hint; повторна та сама помилка → targeted_hint."),
        ("Режим пояснення", "модульні тести, контрольна перевірка (npm run check), Playwright", "пройдено", "Підтверджено, що кнопка пояснення з нейтральним запитом Що робити далі? переводить систему в concept_explanation, а не спирається лише на текстові евристики."),
        ("Валідація структури JSON", "модульні тести (npm test)", "пройдено", "Перевірено null, масив, рядок і частковий об'єкт для JSON-пакетів вправи та фідбеку."),
        ("Sandbox worker", "модульні тести (npm test)", "пройдено", "Перевірено missing function, timeout path, functionName validation і обробку помилок worker."),
        ("Експорт сесії", "модульні тести, контрольна перевірка (npm run check), Playwright", "пройдено", "Підтверджено поля topic, difficulty, attemptsCount, firstFailure, lastAction, transcript, finalRunStatus, timestamp; firstFailure тепер належить поточній вправі."),
        ("Швидка контрольна перевірка", "контрольна перевірка (npm run check)", "пройдено", "Скрипт без важкого стеку швидко перевіряє потік пояснення та формування JSON-експорту сесії."),
        ("Повний сценарій з локальною Ollama", "Playwright + локальна Ollama", "пройдено", "Підтверджено: генерація вправи з джерелом ollama, дві послідовні невдалі спроби з ескалацією підказки, явний режим пояснення та завантаження JSON-експорту."),
    ]
    for row_index, row_values in enumerate(checks, start=1):
        for column_index, value in enumerate(row_values):
            table.cell(row_index, column_index).text = value

    for table in document.tables[1:]:
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
                header = OxmlElement("w:tblHeader")
                header.set(qn("w:val"), "true")
                tr_pr.append(header)

            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.style = "Normal"
                    paragraph.paragraph_format.left_indent = Mm(0)
                    paragraph.paragraph_format.first_line_indent = Mm(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_run_font(run, size=12, bold=row_index == 0)
                    set_paragraph_language(paragraph, UK_LANG)


def add_listing_caption(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, style_name="ListingCaption")
    set_paragraph_text(paragraph, text, style_name="ListingCaption", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    return paragraph


def add_source_note(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, style_name="SourceNote")
    set_paragraph_text(paragraph, text, style_name="SourceNote", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    return paragraph


def add_prompt_or_code_block(anchor: Paragraph, text: str, style_name: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, style_name=style_name)
    size = 10.5 if style_name == "PromptBlock" else 9.5
    set_paragraph_text(
        paragraph,
        text,
        style_name=style_name,
        size=size,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font_name="Courier New",
    )
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    shade_paragraph(paragraph, "F2F2F2")
    return paragraph


def add_dialogue_paragraph(anchor: Paragraph, speaker: str, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, style_name="DialogueBlock")
    clear_paragraph(paragraph)
    paragraph.style = "DialogueBlock"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Mm(-10)
    paragraph.paragraph_format.left_indent = Mm(10)
    label = paragraph.add_run(f"{speaker}: ")
    set_run_font(label, size=14, bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, size=14)
    set_paragraph_language(paragraph, UK_LANG)
    return paragraph


def insert_picture_after(anchor: Paragraph, image_path: Path, width_mm: float) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, style_name="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Mm(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))
    set_paragraph_language(paragraph, UK_LANG)
    return paragraph


def rebuild_appendices(document: Document) -> None:
    appendix_a = find_paragraph_exact(document, "Додаток А. Діаграми архітектури системи")
    appendix_b = find_paragraph_exact(document, "Додаток Б. Приклади промптів")
    appendix_v = find_paragraph_exact(document, "Додаток В. Приклади діалогів агента")
    appendix_g = find_paragraph_exact(document, "Додаток Г. Фрагменти програмного коду")

    remove_content_between(appendix_a, appendix_b)
    remove_content_between(appendix_b, appendix_v)
    remove_content_between(appendix_v, appendix_g)

    end_node = appendix_g._p.getnext()
    while end_node is not None:
        next_node = end_node.getnext()
        remove_element(end_node)
        end_node = next_node

    cursor = appendix_a
    intro = insert_paragraph_after(cursor, style_name="Normal")
    intro.text = (
        "У додатку наведено зменшені копії ключових діаграм, використаних у роботі. Вони "
        "допомагають швидко звірити архітектурні рішення без повернення до основних розділів."
    )
    cursor = intro

    appendix_images = [
        ("Діаграма А.1. Копія рисунка 3.1: компонентна архітектура АІНС.", DIAGRAM_DIR / "figure_3_1_component_architecture.png"),
        ("Діаграма А.2. Копія рисунка 3.2: структура UserState та зв'язок із SessionMemory.", DIAGRAM_DIR / "figure_3_2_session_memory.png"),
        ("Діаграма А.3. Копія рисунка 3.3: послідовність сценарію навчальної взаємодії.", DIAGRAM_DIR / "figure_3_3_sequence.png"),
        ("Діаграма А.4. Копія рисунка 4.1: схема виконання прототипу з локальною LLM.", DIAGRAM_DIR / "figure_4_1_runtime.png"),
    ]
    for label, image_path in appendix_images:
        cursor = add_listing_caption(cursor, label)
        cursor = insert_picture_after(cursor, image_path, width_mm=138)
        cursor = add_source_note(cursor, "Примітка: зменшена копія рисунка з основного тексту роботи.")

    cursor = appendix_b
    intro = insert_paragraph_after(cursor, style_name="Normal")
    intro.text = (
        "Нижче подано скорочені, але автентичні фрагменти промптів, що реально використовуються "
        "у прототипі для генерації вправ і формування адаптивного фідбеку."
    )
    cursor = intro

    cursor = add_listing_caption(cursor, "Лістинг Б.1. Фрагмент промпту для генерації вправи.")
    cursor = add_source_note(cursor, "Джерело: prototype/src/services/exerciseGenerator.js")
    cursor = add_prompt_or_code_block(
        cursor,
        "Ти виконуєш роль репетитора з базового JavaScript.\n"
        "Поверни лише JSON-об'єкт без пояснень.\n\n"
        "Вимоги:\n"
        "- одна невелика вправа для початківця;\n"
        "- не вимагай DOM, мережу, файли, зовнішні бібліотеки;\n"
        "- очікується одна функція;\n"
        "- не включай повний розв'язок.",
        "PromptBlock",
    )

    cursor = add_listing_caption(cursor, "Лістинг Б.2. Фрагмент промпту для адаптивного фідбеку.")
    cursor = add_source_note(cursor, "Джерело: prototype/src/services/feedbackEvaluator.js")
    cursor = add_prompt_or_code_block(
        cursor,
        "Ти - доброзичливий репетитор з JavaScript для початківця.\n"
        "Поверни лише JSON-об'єкт без пояснювального тексту поза JSON.\n\n"
        "Обмеження:\n"
        "- не давай повний розв'язок;\n"
        "- дай короткий адаптивний фідбек українською;\n"
        "- врахуй attemptsCount і тип дії policy.action;\n"
        "- якщо студент просить пояснення, дай концептуальне пояснення без готового коду.",
        "PromptBlock",
    )

    cursor = appendix_v
    intro = insert_paragraph_after(cursor, style_name="Normal")
    intro.text = "Приклади нижче показують форму відповіді агента в типових ситуаціях без розтягування тексту по ширині сторінки."
    cursor = intro

    cursor = add_listing_caption(cursor, "Приклад В.1. Ситуація першої помилки.")
    cursor = add_dialogue_paragraph(cursor, "Студент", "Я запустив тести, але один сценарій не проходить.")
    cursor = add_dialogue_paragraph(cursor, "Репетитор", "Подивись на найпростіший приклад для порожнього масиву.")
    cursor = add_dialogue_paragraph(cursor, "Репетитор", "Не переписуй усе рішення, спочатку перевір, яке значення функція має повернути в цьому випадку.")

    cursor = add_listing_caption(cursor, "Приклад В.2. Запит на пояснення.")
    cursor = add_dialogue_paragraph(cursor, "Студент", "Поясни, чому тут потрібен цикл.")
    cursor = add_dialogue_paragraph(cursor, "Репетитор", "Цикл потрібен, коли треба пройти всі елементи масиву і поступово оновлювати результат.")
    cursor = add_dialogue_paragraph(cursor, "Репетитор", "Спробуй спершу сформулювати, яка змінна має накопичувати суму.")

    cursor = appendix_g
    intro = insert_paragraph_after(cursor, style_name="Normal")
    intro.text = (
        "Нижче наведено короткі фрагменти коду з репозиторію, які ілюструють ключові рішення "
        "політики репетитора, безпечного запуску тестів і експорту стану сесії."
    )
    cursor = intro

    cursor = add_listing_caption(cursor, "Лістинг Г.1. Явний режим пояснення в модулі TutorPolicy.")
    cursor = add_source_note(cursor, "Джерело: prototype/src/services/tutorPolicy.js")
    cursor = add_prompt_or_code_block(
        cursor,
        "if (hasExplicitExplanationMode({ interactionMode, forceExplanation })) {\n"
        "  return {\n"
        "    action: \"concept_explanation\",\n"
        "    hintLevel: \"conceptual\",\n"
        "    rationale: \"Увімкнено явний режим пояснення.\"\n"
        "  };\n"
        "}",
        "CodeBlock",
    )

    cursor = add_listing_caption(cursor, "Лістинг Г.2. Перевірка sandbox-запиту перед виконанням тестів.")
    cursor = add_source_note(cursor, "Джерело: prototype/src/services/sandboxExecution.js")
    cursor = add_prompt_or_code_block(
        cursor,
        "export async function executeSandboxRequest(payload) {\n"
        "  const totalCount = Array.isArray(payload?.tests) ? payload.tests.length : 0;\n"
        "  const validation = validateSandboxRequest(payload);\n\n"
        "  if (!validation.ok) {\n"
        "    return buildSandboxErrorResult({\n"
        "      totalCount,\n"
        "      syntaxError: `Некоректний sandbox payload: ${formatValidationErrors(validation)}`\n"
        "    });\n"
        "  }\n"
        "}",
        "CodeBlock",
    )

    cursor = add_listing_caption(cursor, "Лістинг Г.3. Формування JSON-експорту поточної сесії.")
    cursor = add_source_note(cursor, "Джерело: prototype/src/state/sessionMemory.js")
    add_prompt_or_code_block(
        cursor,
        "buildSessionExport({ timestamp = new Date().toISOString() } = {}) {\n"
        "  return {\n"
        "    topic: this.exercise?.topic ?? this.userState.currentTopic,\n"
        "    difficulty: this.exercise?.difficulty ?? null,\n"
        "    attemptsCount: this.userState.attemptsCount,\n"
        "    firstFailure: this.firstFailure ? structuredClone(this.firstFailure) : null,\n"
        "    finalRunStatus: this.lastRunResult\n"
        "      ? { status: this.lastRunResult.status }\n"
        "      : null,\n"
        "    timestamp\n"
        "  };\n"
        "}",
        "CodeBlock",
    )


def dominant_language(text: str) -> str:
    cyrillic = sum(1 for char in text if "А" <= char <= "я" or char in "ЇїІіЄєҐґ")
    latin = sum(1 for char in text if "A" <= char <= "z")
    return EN_LANG if latin > cyrillic and cyrillic == 0 else UK_LANG


def apply_languages(document: Document) -> None:
    in_abstract = False
    in_references = False

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "ABSTRACT":
            in_abstract = True
            in_references = False
        elif text == "ЗМІСТ":
            in_abstract = False
        elif text == "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            in_references = True
        elif text == "ДОДАТКИ":
            in_references = False

        if paragraph.style.name in {"PromptBlock", "CodeBlock"}:
            set_paragraph_language(paragraph, EN_LANG)
        elif in_abstract:
            set_paragraph_language(paragraph, EN_LANG)
        elif in_references:
            set_paragraph_language(paragraph, dominant_language(text))
        else:
            set_paragraph_language(paragraph, UK_LANG)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_language(paragraph, dominant_language(paragraph.text))


def normalize_spaces(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def heading_variants(text: str) -> set[str]:
    normalized = normalize_spaces(text)
    return {
        normalized,
        normalize_spaces(normalized.replace("→", "->")),
        normalize_spaces(normalized.replace("->", "→")),
    }


def render_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation=file://{(pdf_path.parent / 'lo_profile').resolve()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if generated != pdf_path:
        generated.replace(pdf_path)


def collect_page_texts(pdf_path: Path) -> list[str]:
    with fitz.open(pdf_path) as pdf:
        return [normalize_spaces(page.get_text("text")) for page in pdf]


def find_heading_pages(pdf_path: Path) -> dict[str, int]:
    page_texts = collect_page_texts(pdf_path)
    pages: dict[str, int] = {}

    early_headings = TOC_LINES[:3]
    later_headings = TOC_LINES[3:]

    cursor = 0
    for heading in early_headings:
        variants = heading_variants(heading)
        for page_index in range(cursor, len(page_texts)):
            if any(variant in page_texts[page_index] for variant in variants):
                pages[heading] = page_index + 1
                cursor = page_index
                break

    cursor = 6
    for heading in later_headings:
        variants = heading_variants(heading)
        for page_index in range(cursor, len(page_texts)):
            if any(variant in page_texts[page_index] for variant in variants):
                pages[heading] = page_index + 1
                cursor = page_index
                break

    missing = [heading for heading in TOC_LINES if heading not in pages]
    if missing:
        raise ValueError(f"Headings not found in rendered PDF: {missing}")
    return pages


def count_figures(document: Document) -> int:
    return sum(1 for paragraph in document.paragraphs if paragraph.text.strip().startswith("Рисунок "))


def count_tables(document: Document) -> int:
    return sum(1 for paragraph in document.paragraphs if paragraph.text.strip().startswith("Таблиця "))


def count_appendices(document: Document) -> int:
    return sum(
        1
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 2" and paragraph.text.strip().startswith("Додаток ")
    )


def count_references(document: Document) -> int:
    in_references = False
    count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            in_references = True
            continue
        if text == "ДОДАТКИ":
            break
        if in_references and text.startswith("["):
            count += 1
    return count


def update_counts_and_toc(document: Document, pdf_path: Path) -> dict[str, int]:
    page_count = len(PdfReader(str(pdf_path)).pages)
    figure_count = count_figures(document)
    table_count = count_tables(document)
    appendix_count = count_appendices(document)
    reference_count = count_references(document)
    heading_pages = find_heading_pages(pdf_path)

    find_paragraph_by_prefix(document, "Магістерська курсова робота:").text = (
        f"Магістерська курсова робота: {page_count} с., {figure_count} рис., {table_count} табл., "
        f"{appendix_count} дод., {reference_count} джерел."
    )
    find_paragraph_by_prefix(document, "Master's coursework project:").text = (
        f"Master's coursework project: {page_count} pages, {figure_count} figures, {table_count} tables, "
        f"{appendix_count} appendices, and {reference_count} references."
    )

    toc_heading = find_paragraph_exact(document, "ЗМІСТ")
    next_heading = find_paragraph_exact(document, "СПИСОК УМОВНИХ ПОЗНАЧЕНЬ")
    toc_paragraphs: list[Paragraph] = []
    node = toc_heading._p.getnext()
    while node is not None and node != next_heading._p:
        if node.tag.endswith("}p"):
            toc_paragraphs.append(Paragraph(node, document))
        node = node.getnext()

    for paragraph, toc_line in zip(toc_paragraphs, TOC_LINES):
        paragraph.text = f"{toc_line}\t{heading_pages[toc_line]}"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Mm(0)
        set_paragraph_language(paragraph, EN_LANG if toc_line == "ABSTRACT" else UK_LANG)

    return {
        "pages": page_count,
        "figures": figure_count,
        "tables": table_count,
        "appendices": appendix_count,
        "references": reference_count,
    }


def export_pngs(pdf_path: Path) -> None:
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    with fitz.open(pdf_path) as pdf:
        for page_index in range(len(pdf)):
            pix = pdf.load_page(page_index).get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
            pix.save(PNG_DIR / f"page_{page_index + 1:02d}.png")


def verify_document(document: Document) -> None:
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    forbidden = ["інформаційних фейків", "дезінформації в соціальних мережах", "СЛУЖБОВІ СТОРІНКИ / ЗАВДАННЯ"]
    for marker in forbidden:
        if marker in full_text:
            raise ValueError(f"Forbidden fragment remained in document: {marker}")


def build_stage_document() -> Document:
    document = Document(str(SOURCE_DOC))
    configure_styles(document)
    prepare_toc_region(document)
    replace_text_content(document)
    format_title_table(document)
    update_tables(document)
    rebuild_appendices(document)
    normalize_paragraphs(document)
    apply_languages(document)
    verify_document(document)
    return document


def main() -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)

    document = build_stage_document()
    document.save(str(STAGE_DOC))
    render_pdf(STAGE_DOC, PDF_PATH)

    shutil.copyfile(STAGE_DOC, TARGET_DOC)
    for _ in range(2):
        document = Document(str(TARGET_DOC))
        update_counts_and_toc(document, PDF_PATH)
        normalize_paragraphs(document)
        apply_languages(document)
        verify_document(document)
        document.save(str(TARGET_DOC))
        render_pdf(TARGET_DOC, PDF_PATH)

    export_pngs(PDF_PATH)


if __name__ == "__main__":
    main()
