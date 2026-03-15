from __future__ import annotations

import subprocess
import shutil
from pathlib import Path

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "output" / "coursework_draft_ua_final_polished.docx"
TARGET_DOC = ROOT / "output" / "coursework_draft_ua_submission_ready.docx"
WORK_DIR = ROOT / "tmp" / "docs" / "submission_ready"
STAGE_DOC = WORK_DIR / "coursework_submission_stage.docx"
PDF_PATH = WORK_DIR / "coursework_draft_ua_submission_ready.pdf"

UK_LANG = "uk-UA"
EN_LANG = "en-US"
NARRATIVE_STYLE_NAMES = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "List Bullet", "List Number"]
TOC_LINES = [
    "СЛУЖБОВІ СТОРІНКИ / ЗАВДАННЯ",
    "АНОТАЦІЯ",
    "ABSTRACT",
    "СПИСОК УМОВНИХ ПОЗНАЧЕНЬ",
    "ВСТУП",
    "РОЗДІЛ 1. АНАЛІЗ ПРЕДМЕТНОЇ ГАЛУЗІ АДАПТИВНИХ ІНТЕЛЕКТУАЛЬНИХ НАВЧАЛЬНИХ СИСТЕМ",
    "1.1. Поняття та класифікація інтелектуальних навчальних систем (ITS/АІНС)",
    "1.2. Архітектурні компоненти АІНС",
    "1.3. Діалогові агенти у навчальних системах",
    "1.4. Огляд сучасних систем-репетиторів на основі ШІ",
    "1.5. Висновки до розділу 1",
    "РОЗДІЛ 2. ВЕЛИКІ МОВНІ МОДЕЛІ ЯК ОСНОВА ДІАЛОГОВОГО АГЕНТА",
    "2.1. Архітектура трансформерів та принципи роботи LLM",
    "2.2. Порівняння LLM для навчальних застосунків",
    "2.3. Проєктування промптів, RAG і донавчання як методи адаптації",
    "2.4. Обмеження LLM: галюцинації, пам'ять, етика, приватність",
    "2.5. Обґрунтування вибору локальної відкритої Llama для прототипу",
    "РОЗДІЛ 3. ПРОЄКТУВАННЯ АРХІТЕКТУРИ АІНС",
    "3.1. Функціональні та нефункціональні вимоги",
    "3.2. Загальна архітектура системи",
    "3.3. Модель студента",
    "3.4. Модель знань і навчального контенту",
    "3.5. Модель і стратегії репетитора",
    "3.5.1. Формальне визначення агента-репетитора",
    "3.5.2. Архітектура агента - компонентна модель",
    "3.5.3. Матриця рішень і стратегія мінімальної підказки",
    "3.5.4. Типові сценарії взаємодії",
    "3.5.5. Цикл роботи агента: міркування → дія → спостереження",
    "3.5.6. Організація пам'яті агента",
    "3.5.7. Роль LLM і її обмеження",
    "РОЗДІЛ 4. РЕАЛІЗАЦІЯ ТА ЕКСПЕРИМЕНТАЛЬНА ПЕРЕВІРКА",
    "4.1. Вибір технологічного стеку",
    "4.2. Реалізація ключових модулів системи",
    "4.3. Проєктування промптів для генерації вправ і фідбеку",
    "4.4. Дизайн та проведення експерименту",
    "4.4.1. Функціональна валідація прототипу",
    "4.4.2. Дизайн педагогічного експерименту",
    "4.5. Аналіз результатів",
    "ВИСНОВКИ",
    "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ",
    "ДОДАТКИ",
    "Додаток А. Діаграми архітектури системи",
    "Додаток Б. Приклади промптів",
    "Додаток В. Приклади діалогів агента",
    "Додаток Г. Фрагменти програмного коду",
]


def set_run_font(run, size: float = 14, bold: bool = False, italic: bool = False, font_name: str = "Times New Roman") -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_language(element, lang: str) -> None:
    lang_el = element.find(qn("w:lang"))
    if lang_el is None:
        lang_el = OxmlElement("w:lang")
        element.append(lang_el)
    lang_el.set(qn("w:val"), lang)
    lang_el.set(qn("w:eastAsia"), lang)
    lang_el.set(qn("w:bidi"), lang)


def set_paragraph_language(paragraph: Paragraph, lang: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_r_pr = p_pr.find(qn("w:rPr"))
    if p_r_pr is None:
        p_r_pr = OxmlElement("w:rPr")
        p_pr.append(p_r_pr)
    set_language(p_r_pr, lang)
    for run in paragraph.runs:
        r_pr = run._element.get_or_add_rPr()
        set_language(r_pr, lang)


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("pPr"):
            continue
        paragraph._p.remove(child)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    style_name: str | None = None,
    bold: bool = False,
    italic: bool = False,
    size: float = 14,
    align: WD_ALIGN_PARAGRAPH | None = None,
    lang: str = UK_LANG,
) -> None:
    clear_paragraph(paragraph)
    if style_name is not None:
        paragraph.style = style_name
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    if align is not None:
        paragraph.alignment = align
    set_paragraph_language(paragraph, lang)


def add_mixed_paragraph(paragraph: Paragraph, parts: list[tuple[str, bool]], lang: str = UK_LANG) -> None:
    clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        run = paragraph.add_run(text)
        set_run_font(run, bold=bold)
    set_paragraph_language(paragraph, lang)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = style_name
    if text:
        run = new_paragraph.add_run(text)
        set_run_font(run)
    return new_paragraph


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def remove_numbering(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def remove_table(table: Table) -> None:
    parent = table._element.getparent()
    if parent is not None:
        parent.remove(table._element)


def iter_block_items(document: Document):
    for child in document.element.body:
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def find_paragraph_by_prefix(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def find_paragraph_by_prefix_and_style(document: Document, prefix: str, style_name: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.style.name == style_name and paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix} [{style_name}]")


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(215.9)
    section.page_height = Mm(279.4)
    section.left_margin = Mm(20)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    for style_name in NARRATIVE_STYLE_NAMES:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        set_language(style._element.get_or_add_rPr(), UK_LANG)

    normal = document.styles["Normal"]
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Mm(12.5)


def normalize_paragraphs(document: Document) -> None:
    in_abstract = False
    in_references = False

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name
        text = paragraph.text
        formula_line = text.startswith(("S = ", "A = {", "T: ", "P: ", "R: ")) or text.lstrip().startswith(
            "celebrate_and_reflect"
        )

        if text == "ABSTRACT":
            in_abstract = True
            in_references = False
        elif text == "ЗМІСТ":
            in_abstract = False
        elif text == "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            in_references = True
        elif text == "ДОДАТКИ":
            in_references = False

        if style_name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.keep_with_next = True
        elif style_name in {"Heading 2", "Heading 3"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.keep_with_next = True
        else:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.keep_with_next = False

        if paragraph.text.startswith("Таблиця "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.keep_with_next = True
        elif paragraph.text.startswith("Рисунок "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.keep_with_next = False
        elif paragraph.text == "ЗМІСТ":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
        elif paragraph.text in TOC_LINES or "\t" in paragraph.text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
        elif formula_line:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.line_spacing = 1.0
        elif style_name == "Normal":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if paragraph.text.strip():
                paragraph.paragraph_format.first_line_indent = Mm(12.5)

        for run in paragraph.runs:
            font_name = "Courier New" if run.font.name == "Courier New" else "Times New Roman"
            set_run_font(
                run,
                size=10.5 if font_name == "Courier New" else 12 if formula_line else 14,
                bold=bool(run.bold),
                italic=bool(run.italic),
                font_name=font_name,
            )

        if in_abstract or in_references:
            set_paragraph_language(paragraph, EN_LANG)
        else:
            set_paragraph_language(paragraph, UK_LANG)


def normalize_tables(document: Document) -> None:
    for table in document.tables:
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            if row_idx == 0 and tr_pr.find(qn("w:tblHeader")) is None:
                header = OxmlElement("w:tblHeader")
                header.set(qn("w:val"), "true")
                tr_pr.append(header)

            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    remove_numbering(paragraph)
                    paragraph.style = "Normal"
                    paragraph.paragraph_format.left_indent = Mm(0)
                    paragraph.paragraph_format.first_line_indent = Mm(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    for run in paragraph.runs:
                        set_run_font(run, size=12, bold=row_idx == 0)
                    set_paragraph_language(paragraph, UK_LANG)


def apply_content_updates(document: Document) -> None:
    find_paragraph_by_prefix(
        document,
        "Предметною областю цієї курсової роботи обрано навчання базового JavaScript."
    ).text = (
        "Предметною областю цієї курсової роботи обрано навчання базового JavaScript. "
        "Такий вибір зручний з методичної та технічної точок зору: для невеликих програмістських "
        "вправ легко сформулювати тестовані критерії коректності, а процес взаємодії студента з "
        "репетитором природно вкладається у цикл «згенерувати задачу → написати розв'язок → "
        "запустити тести → отримати адаптивний фідбек» [14], [15], [16], [17], [18]."
    )
    set_paragraph_language(
        find_paragraph_by_prefix(document, "Предметною областю цієї курсової роботи обрано навчання"),
        UK_LANG,
    )

    find_paragraph_by_prefix(
        document,
        "Архітектурно knowledge memory поділяється на чотири типи сутностей:"
    ).text = (
        "Архітектурно пам'ять знань поділяється на чотири типи сутностей: концепти JavaScript, "
        "шаблони задач, перевірочні тести та пояснювальні матеріали. У повній версії системи ці "
        "сутності можуть бути пов'язані у вигляді графа концептів і механізму вибірки релевантних "
        "фрагментів; у поточній версії їх роль частково виконує вбудований резервний набір вправ і "
        "структурований промпт-контракт для LLM."
    )

    formula_paragraph = find_paragraph_by_prefix(document, "S = K × E × C × Topic × Attempts × M")
    set_paragraph_text(
        formula_paragraph,
        "S = K × E × C × Topic × Attempts × M",
        style_name="Normal",
        size=12,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    formula_paragraph.paragraph_format.first_line_indent = Mm(0)
    formula_paragraph.paragraph_format.line_spacing = 1.0

    current = formula_paragraph
    for line in [
        "A = {generate_exercise, give_minimal_hint, give_targeted_hint, explain_concept,",
        "      celebrate_and_reflect, request_retry, escalate}",
        "T: S × O → S",
        "P: S × O → A",
        "R: S × A × Ctx → Text",
    ]:
        current = insert_paragraph_after(current, style_name="Normal")
        set_paragraph_text(
            current,
            line,
            style_name="Normal",
            size=12,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        current.paragraph_format.first_line_indent = Mm(0)
        current.paragraph_format.line_spacing = 1.0

    explanation_paragraph = find_paragraph_by_prefix(document, "Тут `K` позначає оцінку знань")
    explanation_paragraph.text = (
        "Тут: `K` - інтегрована оцінка засвоєння теми; `E` - історія типових помилок; `C` - "
        "робоча оцінка впевненості; `Topic` - поточна тема; `Attempts` - кількість спроб; `M` - "
        "пам'ять сесії; `O` - нові спостереження; `A` - множина педагогічних дій; `T` - функція "
        "переходу стану; `P` - політика вибору дії; `R` - функція формування відповіді; `Ctx` - "
        "контекст генерації; `Text` - текстова відповідь агента. Такий запис відокремлює стан, "
        "політику та мовну генерацію і робить архітектуру агента придатною для формального опису "
        "без ототожнення її з однією лише LLM [11], [12]."
    )

    find_paragraph_by_prefix(
        document,
        "У коді ця матриця реалізована в модулі `TutorPolicy` через набір явних правил."
    ).text = (
        "У коді ця матриця реалізована в модулі `TutorPolicy` через набір явних правил. Для більш "
        "складної системи вона може бути замінена рушієм правил або адаптивною стратегією вибору "
        "дії, але для курсової роботи важливо насамперед показати сам принцип явного розділення "
        "педагогічної дії та генеративної реалізації."
    )

    bullet_updates = {
        "Правильна відповідь:": "Правильна відповідь. ",
        "Помилка:": "Перша помилка. ",
        "Повторна помилка:": "Повторна помилка. ",
        "Запит на пояснення:": "Запит на пояснення. ",
    }
    for prefix, lead in bullet_updates.items():
        paragraph = find_paragraph_by_prefix(document, prefix)
        tail = paragraph.text.split(":", 1)[1].strip()
        if tail:
            tail = f"{tail[0].upper()}{tail[1:]}"
        remove_numbering(paragraph)
        add_mixed_paragraph(paragraph, [(lead, True), (tail, False)])

    for prefix, replacement in {
        "Робочий цикл агента описується схемою «міркування -> дія -> спостереження»":
            "Робочий цикл агента описується схемою «міркування → дія → спостереження», яка "
            "відповідає підходу ReAct. На етапі міркування система інтерпретує нові "
            "спостереження й оновлює `UserState`. На етапі дії обирається педагогічна дія: "
            "створити вправу, дати підказку, пояснити концепт або запросити нову спробу. На "
            "етапі спостереження система отримує нові результати, зокрема результати тестів, і "
            "повторює цикл [12].",
        "Рисунок 3.3 - Діаграма послідовності для сценарію «згенерувати вправу -> пройти тести -> дати фідбек».": 
            "Рисунок 3.3 - Діаграма послідовності для сценарію «згенерувати вправу → пройти тести → дати фідбек».",
    }.items():
        paragraph = find_paragraph_by_prefix(document, prefix)
        paragraph.text = replacement

    heading_355 = find_paragraph_by_prefix_and_style(
        document,
        "3.5.5. Цикл роботи агента: міркування",
        "Heading 3",
    )
    set_paragraph_text(
        heading_355,
        "3.5.5. Цикл роботи агента: міркування → дія → спостереження",
        style_name="Heading 3",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    for paragraph in document.paragraphs:
        if paragraph.text.startswith(
            "Рисунок 3.3 - Діаграма послідовності для сценарію «згенерувати вправу"
        ):
            paragraph.text = (
                "Рисунок 3.3 - Діаграма послідовності для сценарію «згенерувати вправу → "
                "пройти тести → дати фідбек»."
            )

    find_paragraph_by_prefix(
        document,
        "У поточному середовищі локальний виконуваний файл `ollama` не був встановлений"
    ).text = (
        "Під час фінальної функціональної валідації локальний `ollama` був доступний разом із "
        "моделлю `llama3.1:8b`, тому цільове середовище виконання вдалося перевірити без "
        "симуляції. Водночас резервний режим збережено в архітектурі як захист від недоступності "
        "локальної моделі або невалідної JSON-відповіді [25]."
    )

    find_paragraph_by_prefix(
        document,
        "У поточній роботі було виконано лише ту функціональну перевірку"
    ).text = (
        "У поточній роботі виконано лише ту функціональну перевірку, яку можна підтвердити "
        "реальними артефактами середовища: `npm test`, `npm run check`, браузерний прогін через "
        "Playwright і реальний локальний виклик `Ollama + llama3.1:8b`. Освітні результати, "
        "UX-оцінки та ефективність навчання тут не вигадуються. Натомість підтверджено коректну "
        "ескалацію підказок, явний режим пояснення, генерацію вправи локальною моделлю, "
        "полегшений JSON-експорт сесії та стійкість прототипу до невалідних LLM-відповідей."
    )

    find_paragraph_by_prefix(
        document,
        "Отже, інженерна функціональність мінімального практичного сценарію підтверджена"
    ).text = (
        "Отже, інженерна функціональність мінімального практичного сценарію підтверджена як "
        "модульними тестами, так і повним браузерним сценарієм із локальною Ollama. Під час "
        "цього прогону було виявлено й усунуто два локальні недоліки: збільшено тайм-аут "
        "LLM-виклику для `llama3.1:8b` і виправлено скидання `firstFailure` під час переходу до "
        "нової вправи."
    )

    find_paragraph_by_prefix(
        document,
        "Реальний педагогічний експеримент у межах цієї курсової роботи ще не проводився"
    ).text = (
        "У межах цієї курсової роботи педагогічний експеримент зі справжніми студентами не "
        "проводився. Тому далі фіксується лише дизайн наступного етапу без емпіричних висновків. "
        "Порівнюваними умовами залишаються АІНС-прототип з адаптивним фідбеком і контрольний "
        "сценарій без адаптивної репетиторської логіки."
    )

    find_paragraph_by_prefix(
        document,
        "Найдоцільнішим для невеликої вибірки видається кросовер-дизайн."
    ).text = (
        "Для невеликої вибірки придатні два споріднені дизайни: класичний A/B або кросовер, у "
        "якому групи змінюють умови після короткого періоду вирівнювання. У цій роботі "
        "пріоритетним вважається кросовер-дизайн, оскільки він зменшує вплив міжгрупових "
        "відмінностей на ранньому етапі дослідження."
    )

    find_paragraph_by_prefix(
        document,
        "Для уникнення спекулятивних висновків у роботі фіксується лише план статистичного аналізу:"
    ).text = (
        "План аналізу обмежується перевіркою розподілу, вибором параметричного або "
        "непараметричного критерію та оцінкою величини ефекту лише після отримання реальних "
        "даних. Жодних вибірок, середніх значень, p-значень чи висновків про педагогічну "
        "ефективність у межах цієї курсової не наводиться."
    )

    block_items = list(iter_block_items(document))
    for index, item in enumerate(block_items):
        if isinstance(item, Paragraph) and item.text.startswith("Таблиця 4.7 - Шаблон таблиці"):
            remove_paragraph(item)
            next_item = block_items[index + 1]
            if isinstance(next_item, Table):
                remove_table(next_item)
            break

    manual_paragraph = find_paragraph_by_prefix(
        document,
        "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ: після проведення реального педагогічного експерименту"
    )
    remove_paragraph(manual_paragraph)

    find_paragraph_by_prefix(
        document,
        "На момент завершення автоматично згенерованої чернетки наявні лише результати інженерної валідації."
    ).text = (
        "На момент завершення роботи наявні лише результати інженерної валідації: модульні тести, "
        "контрольний сценарій `npm run check` і браузерний прогін через Playwright з локальною "
        "`Ollama + llama3.1:8b`."
    )

    find_paragraph_by_prefix(
        document,
        "Практичний результат полягає в тому, що документ і код узгоджено."
    ).text = (
        "Цих артефактів достатньо, щоб стверджувати: у репозиторії є реальний мінімальний "
        "практичний сценарій, який структурно відповідає описаній архітектурі та відпрацьовує "
        "повний цикл взаємодії з локальною моделлю. Зокрема підтверджено генерацію вправи, "
        "ескалацію підказок після повторної помилки, явний режим пояснення та коректний "
        "JSON-експорт сесії."
    )

    find_paragraph_by_prefix(
        document,
        "Ключове обмеження полягає в тому, що локальне середовище виконання Ollama не було встановлене"
    ).text = (
        "Водночас отримані результати не дають підстав робити висновки про педагогічну "
        "ефективність. Для цього потрібні окреме дослідження зі студентами, визначена вибірка, "
        "заздалегідь зафіксований план аналізу та реальні кількісні показники."
    )

    find_paragraph_by_prefix(
        document,
        "Подальше розширення розділу 4 можливе лише після збору реальних освітніх даних"
    ).text = (
        "Отже, розділ 4 завершується інженерною валідацією і проєктом наступного емпіричного "
        "етапу, а не заявами про вже досягнутий освітній ефект."
    )

    validation_table = document.tables[9]
    validation_table.cell(2, 1).text = "`npm test`, `npm run check`, Playwright"
    validation_table.cell(2, 3).text = (
        "Підтверджено, що кнопка пояснення з нейтральним запитом `Що робити далі?` переводить "
        "систему в `concept_explanation`, а не спирається лише на текстові евристики."
    )
    validation_table.cell(5, 1).text = "`npm test`, `npm run check`, Playwright"
    validation_table.cell(5, 3).text = (
        "Підтверджено поля `topic`, `difficulty`, `attemptsCount`, `firstFailure`, `lastAction`, "
        "`transcript`, `finalRunStatus`, `timestamp`; `firstFailure` тепер належить поточній вправі."
    )
    validation_table.cell(7, 1).text = "Playwright + локальна Ollama"
    validation_table.cell(7, 2).text = "пройдено"
    validation_table.cell(7, 3).text = (
        "Підтверджено: генерація вправи з джерелом `ollama`, дві послідовні невдалі спроби з "
        "ескалацією підказки, явний режим пояснення та завантаження JSON-експорту."
    )

    experiment_plan_table = document.tables[11]
    experiment_plan_table.cell(1, 1).text = (
        "кросовер або A/B із попереднім і підсумковим тестом; обидві умови проходять одну й ту "
        "саму тематичну міні-сесію."
    )
    experiment_plan_table.cell(4, 1).text = (
        "самооцінка впевненості, суб'єктивна корисність підказок, потреба в повторному поясненні."
    )
    experiment_plan_table.cell(5, 1).text = (
        "парний t-критерій або критерій Вілкоксона; величину ефекту розраховувати лише після "
        "отримання реальних даних."
    )


def enhance_figures(document: Document) -> None:
    if len(document.inline_shapes) >= 4:
        for index in [2, 3]:
            shape = document.inline_shapes[index]
            if shape.width.inches < 6.8:
                ratio = shape.height / shape.width
                shape.width = Inches(6.8)
                shape.height = int(shape.width * ratio)


def count_references(document: Document) -> int:
    count = 0
    in_references = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            in_references = True
            continue
        if text == "ДОДАТКИ":
            break
        if in_references and text.startswith("["):
            count += 1
    return count


def render_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation=file://{(pdf_path.parent / 'lo_profile').resolve()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if generated != pdf_path:
        generated.replace(pdf_path)


def normalize_spaces(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def heading_variants(text: str) -> set[str]:
    normalized = normalize_spaces(text)
    return {
        normalized,
        normalize_spaces(normalized.replace("→", "->")),
        normalize_spaces(normalized.replace("->", "→")),
    }


def collect_page_texts(pdf_path: Path) -> list[str]:
    with fitz.open(pdf_path) as pdf:
        return [normalize_spaces(page.get_text("text")) for page in pdf]


def find_heading_pages(pdf_path: Path) -> dict[str, int]:
    page_texts = collect_page_texts(pdf_path)
    pages: dict[str, int] = {}

    early_headings = TOC_LINES[:3]
    later_headings = TOC_LINES[3:]

    cursor = 0
    for heading in early_headings:
        for page_index in range(cursor, len(page_texts)):
            variants = heading_variants(heading)
            if any(variant in page_texts[page_index] for variant in variants):
                pages[heading] = page_index + 1
                cursor = page_index
                break

    cursor = 6
    for heading in later_headings:
        for page_index in range(cursor, len(page_texts)):
            variants = heading_variants(heading)
            if any(variant in page_texts[page_index] for variant in variants):
                pages[heading] = page_index + 1
                cursor = page_index
                break

    missing = [heading for heading in TOC_LINES if heading not in pages]
    if missing:
        raise ValueError(f"Headings not found in rendered PDF: {missing}")
    return pages


def update_toc_and_counts(document: Document, pdf_path: Path) -> None:
    page_count = len(PdfReader(str(pdf_path)).pages)
    figure_count = len(document.inline_shapes)
    table_count = len(document.tables)
    appendix_count = sum(
        1 for paragraph in document.paragraphs if paragraph.style.name == "Heading 2" and paragraph.text.startswith("Додаток ")
    )
    reference_count = count_references(document)

    find_paragraph_by_prefix(document, "Магістерська курсова робота:").text = (
        f"Магістерська курсова робота: {page_count} с., {figure_count} рис., {table_count} табл., "
        f"{appendix_count} дод., {reference_count} джерел."
    )
    find_paragraph_by_prefix(document, "Master's coursework project:").text = (
        f"Master's coursework project: {page_count} pages, {figure_count} figures, {table_count} tables, "
        f"{appendix_count} appendices, and {reference_count} references."
    )

    heading_pages = find_heading_pages(pdf_path)
    for toc_line in TOC_LINES:
        paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if "\t" in paragraph.text and heading_variants(paragraph.text.split("\t", 1)[0]) & heading_variants(toc_line)
        )
        paragraph.text = f"{toc_line}\t{heading_pages[toc_line]}"
        paragraph.paragraph_format.first_line_indent = Mm(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_language(paragraph, UK_LANG if toc_line != "ABSTRACT" else EN_LANG)


def write_document(source: Path, destination: Path) -> None:
    document = Document(str(source))
    configure_styles(document)
    apply_content_updates(document)
    enhance_figures(document)
    normalize_paragraphs(document)
    normalize_tables(document)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))


def finalize_document(target_doc: Path, pdf_path: Path) -> None:
    document = Document(str(target_doc))
    update_toc_and_counts(document, pdf_path)
    normalize_paragraphs(document)
    normalize_tables(document)
    document.save(str(target_doc))


def main() -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    write_document(SOURCE_DOC, STAGE_DOC)
    render_pdf(STAGE_DOC, PDF_PATH)

    shutil.copyfile(STAGE_DOC, TARGET_DOC)
    finalize_document(TARGET_DOC, PDF_PATH)
    render_pdf(TARGET_DOC, PDF_PATH)
    finalize_document(TARGET_DOC, PDF_PATH)
    render_pdf(TARGET_DOC, PDF_PATH)

    with fitz.open(PDF_PATH) as pdf:
        for page_number in [0, 2, 3, 20, 22, 29, 30, 31, 40, len(pdf) - 1]:
            if 0 <= page_number < len(pdf):
                page = pdf[page_number]
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                pix.save(WORK_DIR / f"page-{page_number + 1}.png")


if __name__ == "__main__":
    main()
