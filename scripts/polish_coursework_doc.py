# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "output" / "coursework_draft_ua_fixed.docx"
TARGET_DOC = ROOT / "output" / "coursework_draft_ua_final_polished.docx"


def replace_paragraph(paragraphs, index: int, new_text: str, expected_prefix: str | None = None) -> None:
  paragraph = paragraphs[index]
  current_text = paragraph.text

  if expected_prefix and not current_text.startswith(expected_prefix):
    raise ValueError(f"Paragraph {index} does not match expected prefix: {expected_prefix!r}")

  paragraph.text = new_text


def set_repeat_table_header(row) -> None:
  tr_pr = row._tr.get_or_add_trPr()
  if tr_pr.find(qn("w:tblHeader")) is None:
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_keep_with_next(paragraph) -> None:
  paragraph.paragraph_format.keep_with_next = True


def set_cell_text(table, row_idx: int, col_idx: int, text: str) -> None:
  table.cell(row_idx, col_idx).text = text


def append_reference_before(paragraph, text: str, style_name: str) -> None:
  new_paragraph = paragraph.insert_paragraph_before(text)
  new_paragraph.style = style_name


def apply_monospace_block(paragraph) -> None:
  paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
  paragraph.paragraph_format.space_before = Pt(0)
  paragraph.paragraph_format.space_after = Pt(6)
  for run in paragraph.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(10.5)


def main() -> None:
  parser = argparse.ArgumentParser()
  parser.add_argument("--page-count", type=int, default=41)
  args = parser.parse_args()

  document = Document(str(SOURCE_DOC))
  paragraphs = document.paragraphs

  replace_paragraph(
    paragraphs,
    23,
    "Службові сторінки за потреби оформлюються за кафедральним шаблоном; у цій версії залишено лише ті реквізити, які мають бути внесені вручну перед поданням роботи.",
    "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ:"
  )
  replace_paragraph(
    paragraphs,
    7,
    "МАГІСТЕРСЬКА КУРСОВА РОБОТА",
    "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ:"
  )
  replace_paragraph(
    paragraphs,
    26,
    f"Магістерська курсова робота: {args.page_count} с., 4 рис., 14 табл., 4 дод., 30 джерел.",
    "Магістерська курсова робота:"
  )
  replace_paragraph(
    paragraphs,
    30,
    "Мета роботи - спроєктувати узгоджену архітектуру АІНС із LLM-ядром і продемонструвати її працездатність на мінімальному практичному сценарії: генерація вправи, написання розв'язку, запуск тестів та надання адаптивного фідбеку.",
    "Мета роботи -"
  )
  replace_paragraph(
    paragraphs,
    35,
    f"Master's coursework project: {args.page_count} pages, 4 figures, 14 tables, 4 appendices, and 30 references.",
    "Master's coursework project:"
  )
  replace_paragraph(
    paragraphs,
    38,
    "The goal is to design a coherent AINS architecture and support it with a minimal working practical scenario that covers exercise generation, student code editing, client-side test execution, and adaptive feedback.",
    "The goal is to design"
  )
  replace_paragraph(
    paragraphs,
    39,
    "The document separates conceptual architecture from practical implementation. A local Ollama-based runtime, Web Worker sandboxing, JSON retry and backup logic, and an honest experimental design without fabricated educational outcomes are treated as core principles.",
    "The document separates"
  )
  replace_paragraph(
    paragraphs,
    56,
    "2.3. Проєктування промптів, RAG і донавчання як методи адаптації\t14",
    "2.3. Prompt engineering"
  )
  replace_paragraph(
    paragraphs,
    58,
    "2.5. Обґрунтування вибору локальної відкритої Llama для прототипу\t14",
    "2.5. Обґрунтування вибору локальної open-source"
  )
  replace_paragraph(
    paragraphs,
    69,
    "3.5.5. Цикл роботи агента: міркування -> дія -> спостереження\t23",
    "3.5.5. Цикл роботи агента:"
  )
  replace_paragraph(
    paragraphs,
    75,
    "4.3. Проєктування промптів для генерації вправ і фідбеку\t30",
    "4.3. Проєктування промптів"
  )
  replace_paragraph(
    paragraphs,
    76,
    "4.4. Дизайн та проведення експерименту\t30",
    "4.4. Дизайн та проведення експерименту"
  )
  replace_paragraph(
    paragraphs,
    77,
    "4.4.1. Функціональна валідація прототипу\t30",
    "4.4.1. Функціональна валідація прототипу"
  )
  replace_paragraph(
    paragraphs,
    78,
    "4.4.2. Дизайн педагогічного експерименту\t33",
    "4.4.2. Дизайн педагогічного експерименту"
  )
  replace_paragraph(
    paragraphs,
    79,
    "4.5. Аналіз результатів\t35",
    "4.5. Аналіз результатів"
  )
  replace_paragraph(
    paragraphs,
    80,
    "ВИСНОВКИ\t37",
    "ВИСНОВКИ"
  )
  replace_paragraph(
    paragraphs,
    81,
    "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ\t38",
    "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ"
  )
  replace_paragraph(
    paragraphs,
    82,
    "ДОДАТКИ\t42",
    "ДОДАТКИ"
  )
  replace_paragraph(
    paragraphs,
    83,
    "Додаток А. Діаграми архітектури системи\t42",
    "Додаток А. Діаграми архітектури системи"
  )
  replace_paragraph(
    paragraphs,
    84,
    "Додаток Б. Приклади промптів\t42",
    "Додаток Б. Приклади промптів"
  )
  replace_paragraph(
    paragraphs,
    85,
    "Додаток В. Приклади діалогів агента\t42",
    "Додаток В. Приклади діалогів агента"
  )
  replace_paragraph(
    paragraphs,
    86,
    "Додаток Г. Фрагменти програмного коду\t43",
    "Додаток Г. Фрагменти програмного коду"
  )
  replace_paragraph(
    paragraphs,
    111,
    "Предмет дослідження - архітектурні моделі, механізми пам'яті, політики прийняття рішень і програмні засоби побудови АІНС на основі LLM для сценарію репетиторського супроводу під час вивчення JavaScript.",
    "Предмет дослідження -"
  )
  replace_paragraph(
    paragraphs,
    137,
    "Для освітніх систем доступні як закриті провідні моделі, так і відкриті локальні моделі. Закриті сервіси зазвичай демонструють вищу якість міркування й пояснень, однак створюють залежність від зовнішнього API, підвищують ризики для приватності й ускладнюють відтворюваність. Відкриті локальні моделі є слабшими в середньому, але дають більше контролю над витратами, обробкою даних і правилами інтеграції [7], [8], [10], [17], [18], [19], [29].",
    "Для освітніх систем доступні"
  )
  replace_paragraph(
    paragraphs,
    138,
    "Для задачі базового JavaScript-репетитора не потрібна максимальна генеративна потужність, натомість важливі локальність, передбачуваність і можливість працювати з невеликим структурованим JSON-контрактом. Саме тому в роботі обрано локальне середовище виконання `Ollama + llama3.1:8b` як цільове рішення для прототипу [25], [30].",
    "Для задачі базового JavaScript tutor"
  )
  replace_paragraph(
    paragraphs,
    139,
    "2.3. Проєктування промптів, RAG і донавчання як методи адаптації",
    "2.3. Prompt engineering"
  )
  replace_paragraph(
    paragraphs,
    140,
    "На практиці адаптація LLM до навчального домену може здійснюватися щонайменше трьома способами: через проєктування промптів, через генерацію з доповненим пошуком (retrieval-augmented generation, RAG) та через донавчання. Проєктування промптів є найдоступнішим шляхом для курсового проєкту, оскільки дозволяє нав'язати формат JSON, заборону на повне розкриття розв'язку і стиль репетитора без окремого етапу донавчання [11], [12], [25].",
    "На практиці адаптація"
  )
  replace_paragraph(
    paragraphs,
    141,
    "RAG є логічним наступним кроком для розвитку системи, адже дозволяє під'єднати зовнішнє сховище навчального контенту, прикладів і теоретичних пояснень. Донавчання у межах поточної курсової не реалізується через обмеження масштабу, але в архітектурному описі воно розглядається як потенційний засіб стабілізації стилю та якості педагогічного фідбеку.",
    "RAG є логічним"
  )
  replace_paragraph(
    paragraphs,
    144,
    "Звідси випливають архітектурні наслідки для курсової роботи: необхідність контролю формату відповіді, повторних спроб із переходом до резервного режиму при невалідному JSON, окремого механізму політики репетитора, обмеження на розкриття розв'язку та безпечного виконання лише тих інструментальних дій, які не вимагають виконання коду на сервері.",
    "Звідси випливають архітектурні наслідки"
  )
  replace_paragraph(
    paragraphs,
    145,
    "2.5. Обґрунтування вибору локальної відкритої Llama для прототипу",
    "2.5. Обґрунтування вибору локальної open-source"
  )
  replace_paragraph(
    paragraphs,
    146,
    "Вибір локальної Llama через Ollama зумовлено кількома факторами. По-перше, це відповідає вимозі використовувати безкоштовне локальне середовище виконання без звернення до платних хмарних API. По-друге, Ollama надає достатньо простий HTTP API для інтеграції у полегшений прототип. По-третє, модель класу 8B достатня для генерації невеликих задач, коротких підказок та структурованих відповідей у межах базового навчального сценарію [17], [18], [19], [25], [30].",
    "Вибір локальної Llama через Ollama"
  )
  replace_paragraph(
    paragraphs,
    148,
    "Отже, LLM у межах цієї роботи виступає не як автономний педагог, а як гнучкий генеративний компонент, вбудований у ширшу архітектуру АІНС з явною моделлю стану, механізмом пам'яті та політикою ухвалення рішень [7], [9], [12], [25], [30].",
    "Отже, LLM у межах цієї роботи"
  )
  replace_paragraph(
    paragraphs,
    151,
    "Система проєктується не як універсальний освітній портал, а як мінімальний практичний сценарій для репетиторського супроводу під час вивчення базового JavaScript. Це дозволяє сформулювати чіткі вимоги й водночас залишити місце для подальшого масштабування.",
    "Система проектується"
  )
  replace_paragraph(
    paragraphs,
    181,
    "Компонентна модель агента складається з таких блоків. Модуль сприйняття інтерпретує вхідний запит і визначає режим взаємодії. Модуль пам'яті акумулює поточний контекст і довготривалі відомості. Модуль міркування оцінює стан студента та обирає педагогічну дію. Модуль дій викликає генерацію вправи, перевірку тестів або пояснення. Модуль генерації формує кінцеву текстову відповідь на основі політики та контексту.",
    "Компонентна модель агента"
  )
  replace_paragraph(
    paragraphs,
    182,
    "У прототипі ця декомпозиція реалізована не повністю: роль модуля сприйняття виконують явний режим `interactionMode: \"explain\"` для кнопки пояснення та прості текстові евристики для вільного запиту, а роль модуля генерації виконує локальна LLM або резервний механізм. Однак саме така схема використовується як основа повнішої системи в магістерському проєкті.",
    "У прототипі ця декомпозиція"
  )
  replace_paragraph(
    paragraphs,
    187,
    "У коді ця матриця реалізована в модулі `TutorPolicy` через набір явних правил. Для більш складної системи вона може бути замінена комбінацією рушія правил, адаптивної стратегії типу multi-armed bandit або навчання політики, але для курсової роботи важливо насамперед показати сам принцип явного розділення педагогічної дії та генеративної реалізації.",
    "У коді ця матриця"
  )
  replace_paragraph(
    paragraphs,
    191,
    "Повторна помилка: агент переходить до більш адресної підказки (`targeted_hint`) і конкретніше називає, який фрагмент логіки слід перевірити.",
    "Повторна помилка:"
  )
  replace_paragraph(
    paragraphs,
    192,
    "Запит на пояснення: агент переходить у режим концептуального пояснення (`concept_explanation`) і дає теоретичне пояснення без завершеного розв'язку.",
    "Запит на пояснення:"
  )
  replace_paragraph(
    paragraphs,
    193,
    "3.5.5. Цикл роботи агента: міркування -> дія -> спостереження",
    "3.5.5. Цикл роботи агента:"
  )
  replace_paragraph(
    paragraphs,
    194,
    "Робочий цикл агента описується схемою «міркування -> дія -> спостереження», яка відповідає підходу ReAct. На етапі міркування система інтерпретує нові спостереження й оновлює `UserState`. На етапі дії обирається педагогічна дія: створити вправу, дати підказку, пояснити концепт або запросити нову спробу. На етапі спостереження система отримує нові результати, зокрема результати тестів, і повторює цикл [12].",
    "Робочий цикл агента"
  )
  replace_paragraph(
    paragraphs,
    209,
    "У поточному середовищі локальний виконуваний файл `ollama` не був встановлений, тому прототип додатково отримав резервний режим. Це важливий практичний висновок: для дослідницької системи потрібно з самого початку передбачати резервну логіку, а не вважати зовнішнє середовище виконання гарантовано доступним [25].",
    "У поточному середовищі"
  )
  replace_paragraph(
    paragraphs,
    217,
    "Такий обсяг реалізації відповідає принципу пріоритету архітектури: замість великої кількості другорядних функцій свідомо реалізовано лише ті модулі, які демонструють зв'язок між генерацією задачі, моделлю стану, тестовим запуском, режимом пояснення та адаптивним репетиторським фідбеком. На рівні `TutorPolicy` і `main.js` додано явний режим пояснення, тому натискання кнопки `Пояснити тему` завжди переводить агента до `concept_explanation`, а на рівні `SessionMemory` та UI реалізовано полегшений JSON-експорт поточної сесії як допоміжний інструмент для подальших досліджень типу A/B або кросовер-досліджень.",
    "Така реалізація відповідає"
  )
  replace_paragraph(
    paragraphs,
    219,
    "Для прототипу сформовано дві основні стратегії формування промптів: генерація вправи та формування адаптивного фідбеку. В обох випадках система вимагає лише JSON-відповідь, що спрощує валідацію й нормалізацію. Такий підхід узгоджується з рекомендаціями щодо керованого використання LLM у навчанні та програмуванні [9], [12], [15], [16], [17], [18], [26], [27].",
    "Для прототипу сформовано"
  )
  replace_paragraph(
    paragraphs,
    228,
    "У поточній роботі було виконано лише ту функціональну перевірку, яку можна підтвердити реальними артефактами середовища. Освітні результати, UX-оцінки та ефективність навчання тут не вигадуються. Натомість модульними тестами й швидкою контрольною перевіркою підтверджено коректну ескалацію підказок, явний режим пояснення, полегшений JSON-експорт сесії та базову стійкість прототипу до невалідних LLM-відповідей.",
    "У поточній роботі було виконано"
  )
  replace_paragraph(
    paragraphs,
    231,
    "Отже, інженерна функціональність мінімального практичного сценарію підтверджена для чистих модулів, явного режиму пояснення, JSON-експорту сесії та базової статичної доставки; повний інтерактивний сценарій з локальним Ollama має бути додатково перевірений вручну у браузері.",
    "Отже, інженерна функціональність"
  )
  replace_paragraph(
    paragraphs,
    233,
    "Реальний педагогічний експеримент у межах цієї курсової роботи ще не проводився, тому тут подається лише чесно оформлений дизайн дослідження, який може бути використаний на наступному етапі [14], [19], [21], [22], [26], [27]. JSON-експорт сесії розглядається як допоміжний інструмент для накопичення структурованих логів у майбутньому A/B- або кросовер-дослідженні.",
    "Реальний педагогічний експеримент"
  )
  replace_paragraph(
    paragraphs,
    234,
    "Найдоцільнішим для невеликої вибірки видається кросовер-дизайн. Одна група спочатку працює з АІНС-прототипом, друга - з контрольним сценарієм без адаптивного репетиторського фідбеку; після короткого періоду вирівнювання умови міняються місцями.",
    "Найдоцільнішим для невеликої вибірки"
  )
  replace_paragraph(
    paragraphs,
    242,
    "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ: після проведення реального педагогічного експерименту додати вибірку, статистичні результати, величини ефекту та інтерпретацію.]",
    "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ:"
  )
  replace_paragraph(
    paragraphs,
    245,
    "Практичний результат полягає в тому, що документ і код узгоджено. У тексті не заявляються модулі, яких немає у реалізації, а в коді реалізовано саме ті компоненти, які описано в підрозділах 4.1-4.3: інтерфейс користувача, `SessionMemory` з JSON-експортом, `TutorPolicy` з явним режимом пояснення, `ExerciseGenerator`, `ClientTestRunner`, `FeedbackEvaluator` та `OllamaClient`.",
    "Практичний результат полягає"
  )
  replace_paragraph(
    paragraphs,
    246,
    "Ключове обмеження полягає в тому, що локальне середовище виконання Ollama не було встановлене у поточному середовищі. Це означає, що частина практичної логіки перевірялася у резервному режимі, а повна робота з реальною локальною моделлю потребує ручного запуску у браузері.",
    "Ключове обмеження полягає"
  )
  replace_paragraph(
    paragraphs,
    247,
    "Подальше розширення розділу 4 можливе лише після збору реальних освітніх даних у межах A/B- або кросовер-експерименту; до цього моменту аналіз результатів має залишатися описом інженерної валідації та плану педагогічного дослідження.",
    "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ:"
  )
  replace_paragraph(
    paragraphs,
    250,
    "Найважливішим результатом є архітектурна схема агента, в якій виокремлено модуль сприйняття, модуль пам'яті, модуль міркування, модуль дій і модуль генерації, а також формально описано стан, дії, політику й функцію переходу.",
    "Найважливішим результатом"
  )
  replace_paragraph(
    paragraphs,
    251,
    "Для практичної частини реалізовано мінімальний практичний сценарій репетиторського супроводу з JavaScript: генерація вправи, введення коду, запуск тестів у `Web Worker`, адаптивний фідбек, `SessionMemory`, явний режим пояснення та резервна логіка у випадку недоступності або нестабільності локальної LLM.",
    "Для практичної частини"
  )
  replace_paragraph(
    paragraphs,
    253,
    "Отриманий прототип і текстова чернетка можуть слугувати базою для подальшого розширення в повну магістерську роботу з додаванням довготривалої пам'яті, RAG, багатшої моделі навчального контенту та реального оцінювання впливу системи на навчальні результати.",
    "Отриманий прототип"
  )
  replace_paragraph(
    paragraphs,
    279,
    "[25] Ollama. API documentation. 2024. URL: https://docs.ollama.com/api (дата звернення: 14.03.2026).",
    "[25]"
  )
  replace_paragraph(
    paragraphs,
    301,
    "Фрагмент 1. Явний режим пояснення у `TutorPolicy`.",
    "Фрагмент 1."
  )
  replace_paragraph(
    paragraphs,
    302,
    "if (forceExplanation || interactionMode === \"explain\") {\n  return { action: \"concept_explanation\" };\n}\nif (isNaturalLanguageExplanationRequest(studentRequest)) {\n  return { action: \"concept_explanation\" };\n}",
    "const currentSignature"
  )
  replace_paragraph(
    paragraphs,
    305,
    "Фрагмент 3. Формування JSON-експорту поточної сесії.",
    "Фрагмент 3."
  )
  replace_paragraph(
    paragraphs,
    306,
    "return {\n  topic: this.exercise?.topic ?? this.userState.currentTopic,\n  difficulty: this.exercise?.difficulty ?? null,\n  attemptsCount: this.userState.attemptsCount,\n  firstFailure: this.firstFailure ? structuredClone(this.firstFailure) : null,\n  finalRunStatus: this.lastRunResult ? { status: this.lastRunResult.status } : null\n};",
    "this.userState.attemptsCount"
  )

  table_7 = document.tables[7]
  set_cell_text(table_7, 1, 0, "Браузерний інтерфейс")
  set_cell_text(table_7, 2, 0, "Середовище LLM")
  set_cell_text(table_7, 2, 2, "локальне відкрите середовище виконання відповідно до вимог роботи")
  set_cell_text(table_7, 3, 1, "SessionMemory у пам'яті процесу")
  set_cell_text(
    table_7,
    3,
    2,
    "простий і прозорий варіант для мінімального практичного сценарію та майбутнього JSON-експорту сесій"
  )
  set_cell_text(
    table_7,
    4,
    2,
    "без виконання коду на сервері та з ізоляцією обчислень"
  )
  set_cell_text(table_7, 5, 2, "перевірка чистої логіки модулів")

  table_8 = document.tables[8]
  set_cell_text(table_8, 1, 0, "Інтерфейс користувача")
  set_cell_text(
    table_8,
    2,
    2,
    "Зберігає `UserState`, поточну вправу, історію повідомлень, першу зафіксовану помилку та формує полегшений JSON-експорт сесії."
  )
  set_cell_text(
    table_8,
    3,
    2,
    "Політика на основі правил з коректною ескалацією: перша помилка -> `minimal_hint`, повторна та сама -> `targeted_hint`, а кнопка пояснення через `interactionMode: \"explain\"` завжди веде до `concept_explanation`."
  )
  set_cell_text(
    table_8,
    4,
    2,
    "Генерація вправи через LLM, явна перевірка структури та перехід до резервної вправи при невідповідному JSON."
  )
  set_cell_text(
    table_8,
    5,
    2,
    "Клієнтський запуск тестів у `Web Worker` з тайм-аутом, перевіркою вхідних даних, `worker.onerror` і `messageerror`."
  )
  set_cell_text(
    table_8,
    6,
    2,
    "Адаптивний фідбек через LLM або резервний режим; резервне пояснення формує змістовне пояснення концепту."
  )
  set_cell_text(
    table_8,
    7,
    2,
    "HTTP API, перевірка доступності, перевірка структури та повторні спроби з переходом до резервного режиму для JSON-відповідей."
  )

  table_9 = document.tables[9]
  set_cell_text(table_9, 1, 0, "Політика підказок")
  set_cell_text(table_9, 1, 1, "`npm test`")
  set_cell_text(table_9, 1, 2, "пройдено")
  set_cell_text(
    table_9,
    1,
    3,
    "Перевірено: перша помилка -> `minimal_hint`; повторна та сама помилка -> `targeted_hint`."
  )
  set_cell_text(table_9, 2, 0, "Режим пояснення")
  set_cell_text(table_9, 2, 1, "`npm test`, `npm run check`")
  set_cell_text(table_9, 2, 2, "пройдено")
  set_cell_text(
    table_9,
    2,
    3,
    "Підтверджено, що явний `interactionMode: \"explain\"` гарантує перехід до `concept_explanation`, а пояснення лишається змістовним."
  )
  set_cell_text(table_9, 3, 0, "Валідація структури JSON")
  set_cell_text(table_9, 3, 1, "`npm test`")
  set_cell_text(table_9, 3, 2, "пройдено")
  set_cell_text(
    table_9,
    3,
    3,
    "Перевірено `null`, масив, рядок і частковий об'єкт для exercise/feedback JSON."
  )
  set_cell_text(table_9, 4, 0, "Sandbox worker")
  set_cell_text(table_9, 4, 1, "`npm test`")
  set_cell_text(table_9, 4, 2, "пройдено")
  set_cell_text(
    table_9,
    4,
    3,
    "Перевірено `missing function`, timeout path, `functionName` validation, `worker` error handling."
  )
  set_cell_text(table_9, 5, 0, "Експорт сесії")
  set_cell_text(table_9, 5, 1, "`npm test`, `npm run check`")
  set_cell_text(table_9, 5, 2, "пройдено")
  set_cell_text(
    table_9,
    5,
    3,
    "Перевірено поля `topic`, `difficulty`, `attemptsCount`, `firstFailure`, `lastAction`, `transcript`, `finalRunStatus`, `timestamp`."
  )
  set_cell_text(table_9, 6, 0, "Швидка контрольна перевірка")
  set_cell_text(table_9, 6, 1, "`npm run check`")
  set_cell_text(table_9, 6, 2, "пройдено")
  set_cell_text(
    table_9,
    6,
    3,
    "Скрипт без важкого стеку швидко перевіряє потік пояснення та формування JSON-експорту сесії."
  )
  set_cell_text(table_9, 7, 0, "Повний сценарій з локальною Ollama")
  set_cell_text(table_9, 7, 1, "ручний запуск у браузері")
  set_cell_text(table_9, 7, 2, "потребує ручної перевірки")
  set_cell_text(
    table_9,
    7,
    3,
    "Потребує встановленого `ollama` і локальної моделі `llama3.1:8b`."
  )

  for table in document.tables:
    set_repeat_table_header(table.rows[0])

  for index in [185, 196, 207, 212, 215, 229, 235, 237, 239, 240]:
    set_keep_with_next(paragraphs[index])

  for index in [204, 205, 213, 218, 226, 227, 232, 243, 248, 254, 284, 285, 295, 300]:
    paragraphs[index].paragraph_format.keep_with_next = True

  for shape_index in [2, 3]:
    shape = document.inline_shapes[shape_index]
    aspect_ratio = shape.height / shape.width
    shape.width = Inches(6.6)
    shape.height = int(shape.width * aspect_ratio)

  reference_style = paragraphs[283].style.name
  appendices_heading = paragraphs[284]
  append_reference_before(
    appendices_heading,
    "[30] Meta. Llama 3.1 model resources. 2024. URL: https://llama.meta.com/ (дата звернення: 14.03.2026).",
    reference_style
  )

  for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if text.startswith("Рисунок ") or text.startswith("Таблиця "):
      paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

  for index in [294, 295, 302, 304, 306]:
    apply_monospace_block(paragraphs[index])

  TARGET_DOC.parent.mkdir(parents=True, exist_ok=True)
  document.save(str(TARGET_DOC))


if __name__ == "__main__":
  main()
