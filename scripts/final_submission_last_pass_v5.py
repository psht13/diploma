# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import shutil
import subprocess
import time
from pathlib import Path

import fitz
from docx import Document
from PIL import Image, ImageDraw, ImageOps


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "output" / "coursework_draft_ua_submission_ready_v4.docx"
TARGET_DOC = ROOT / "output" / "coursework_draft_ua_submission_ready_v5.docx"
RENDER_DIR = ROOT / "tmp" / "docs" / "submission_ready_v5_final"
PDF_PATH = RENDER_DIR / f"{TARGET_DOC.stem}.pdf"
PNG_DIR = RENDER_DIR / "pages"
PREVIEW_DIR = RENDER_DIR / "previews"

EARLY_FRONTMATTER_TITLES = {"АНОТАЦІЯ", "ABSTRACT"}
FORBIDDEN_TERMS = [
    "manual",
    "вручну",
    "ручне втручання",
    "ручний ввід",
    "не вигадувалося",
    "чесно оформлений",
    "автентичні фрагменти",
    "Playwright",
    "npm",
    "dev-log",
    "workflow",
]


def normalize(text: str) -> str:
    return " ".join(text.split())


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_paragraph_text(doc: Document, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text == old:
            set_paragraph_text(paragraph, new)
            return
    raise ValueError(f"Paragraph not found for replacement: {old}")


def set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    set_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph_text(paragraph, "")


def count_metrics(doc: Document) -> dict[str, int]:
    figure_count = sum(
        1 for p in doc.paragraphs if normalize(p.text).startswith("Рисунок ")
    )
    table_count = sum(
        1 for p in doc.paragraphs if normalize(p.text).startswith("Таблиця ")
    )
    appendix_count = sum(
        1
        for p in doc.paragraphs
        if p.style.name.startswith("Heading")
        and re.match(r"^Додаток [А-ЯІЇЄҐA-Z]\.", normalize(p.text))
    )
    reference_count = sum(
        1 for p in doc.paragraphs if re.match(r"^\[\d+\]", normalize(p.text))
    )
    return {
        "figures": figure_count,
        "tables": table_count,
        "appendices": appendix_count,
        "references": reference_count,
    }


def render_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    profile_dir = Path("/tmp") / f"lo_profile_codex_v5_{os.getpid()}_{int(time.time() * 1000)}"
    profile_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if generated != pdf_path:
        shutil.copy2(generated, pdf_path)


def extract_page_texts(pdf_path: Path) -> tuple[list[str], int]:
    pdf = fitz.open(pdf_path)
    texts = [normalize(page.get_text("text")) for page in pdf]
    page_count = pdf.page_count
    pdf.close()
    return texts, page_count


def split_toc_entry(text: str) -> tuple[str, str, int]:
    match = re.match(r"^(.*?)(\t+|\s+)(\d+)$", text)
    if not match:
        raise ValueError(f"Cannot parse TOC entry: {text!r}")
    title, separator, page = match.groups()
    return title, separator, int(page)


def locate_title_page(title: str, page_texts: list[str]) -> int:
    matches = [index + 1 for index, page_text in enumerate(page_texts) if title in page_text]
    if not matches:
        raise ValueError(f"Title not found in rendered PDF: {title}")
    if title in EARLY_FRONTMATTER_TITLES:
        return min(matches)
    return max(matches)


def sync_counts_and_toc(doc_path: Path, pdf_path: Path) -> tuple[dict[str, int], dict[str, int], bool]:
    page_texts, page_count = extract_page_texts(pdf_path)
    doc = Document(doc_path)
    metrics = count_metrics(doc)
    changed = False

    ua_counts = (
        f"Магістерська курсова робота: {page_count} с., {metrics['figures']} рис., "
        f"{metrics['tables']} табл., {metrics['appendices']} дод., {metrics['references']} джерел."
    )
    en_counts = (
        f"Master's coursework project: {page_count} pages, {metrics['figures']} figures, "
        f"{metrics['tables']} tables, {metrics['appendices']} appendices, and {metrics['references']} references."
    )

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Магістерська курсова робота:"):
            if paragraph.text != ua_counts:
                set_paragraph_text(paragraph, ua_counts)
                changed = True
        elif paragraph.text.startswith("Master's coursework project:"):
            if paragraph.text != en_counts:
                set_paragraph_text(paragraph, en_counts)
                changed = True

    toc_range = range(38, 82)
    page_map: dict[str, int] = {}
    for index in toc_range:
        paragraph = doc.paragraphs[index]
        title, separator, current_page = split_toc_entry(paragraph.text)
        actual_page = locate_title_page(title, page_texts)
        page_map[title] = actual_page
        new_text = f"{title}{separator}{actual_page}"
        if paragraph.text != new_text:
            set_paragraph_text(paragraph, new_text)
            changed = True

    if changed:
        doc.save(doc_path)

    metrics["pages"] = page_count
    return metrics, page_map, changed


def apply_content_edits(doc: Document) -> None:
    replace_paragraph_text(
        doc,
        "У прототипі ця модель навмисно спрощена до JSON-репрезентації вправи з полями title, prompt, starterCode, functionName, tests, concepts, rubric. Таке рішення дозволяє зберегти місце для майбутнього RAG, але не розширювати практичну частину понад необхідний мінімум.",
        "У прототипі ця модель навмисно спрощена до JSON-подання вправи, у якому фіксуються назва вправи, формулювання задачі, початковий код, назва функції, тестові приклади, перелік понять і критерії оцінювання. Таке рішення дозволяє зберегти місце для майбутнього RAG, але не розширювати практичну частину понад необхідний мінімум.",
    )
    replace_paragraph_text(
        doc,
        "S = K × E × C × Topic × Attempts × M",
        "S = K × E × C × Th × N × M",
    )
    replace_paragraph_text(
        doc,
        "R: S × A × Ctx → Text",
        "R: S × A × Cx → V",
    )
    replace_paragraph_text(
        doc,
        "Тут: K - інтегрована оцінка засвоєння теми; E - історія типових помилок; C - робоча оцінка впевненості; Topic - поточна тема; Attempts - кількість спроб; M - пам'ять сесії; O - нові спостереження; A - множина педагогічних дій; T - функція переходу стану; P - політика вибору дії; R - функція формування відповіді; Ctx - контекст генерації; Text - текстова відповідь агента. Такий запис відокремлює стан, політику та мовну генерацію і робить архітектуру агента придатною для формального опису без ототожнення її з однією лише LLM [11], [12].",
        "У цьому записі змінна стану охоплює інтегровану оцінку засвоєння теми (K), історію типових помилок (E), робочу оцінку впевненості (C), поточну тему навчання (Th), кількість спроб (N) і пам'ять сесії (M); O позначає нові спостереження, A - множину педагогічних дій, T - функцію переходу стану, P - політику вибору дії, R - функцію формування відповіді, а Cx і V - відповідно контекст відповіді та сформовану текстову відповідь агента. Такий запис відокремлює стан, політику та мовну генерацію і робить архітектуру агента придатною для формального опису без ототожнення її з однією лише LLM [11], [12].",
    )
    replace_paragraph_text(
        doc,
        "Такий обсяг реалізації відповідає принципу пріоритету архітектури: замість великої кількості другорядних функцій свідомо реалізовано лише ті модулі, які демонструють зв'язок між генерацією задачі, моделлю стану, тестовим запуском, режимом пояснення та адаптивним репетиторським фідбеком. На рівні педагогічної політики та основного клієнтського сценарію додано явний режим пояснення, тому натискання кнопки «Пояснити тему» завжди переводить агента до концептуального пояснення, а на рівні пам'яті сесії та UI реалізовано полегшений JSON-експорт поточної сесії як допоміжний інструмент для подальших досліджень типу A/B або кросовер-досліджень.",
        "Такий обсяг реалізації відповідає принципу пріоритету архітектури: замість великої кількості другорядних функцій свідомо реалізовано лише ті модулі, які демонструють зв'язок між генерацією задачі, моделлю стану, тестовим запуском, режимом пояснення та адаптивним репетиторським фідбеком. На рівні педагогічної політики та основного клієнтського сценарію додано явний режим пояснення, тому натискання кнопки «Пояснити тему» завжди переводить агента до концептуального пояснення, а на рівні пам'яті сесії та інтерфейсу користувача реалізовано експорт даних поточної сесії у форматі JSON як допоміжний інструмент для подальших досліджень типу A/B або кросовер-досліджень.",
    )
    replace_paragraph_text(
        doc,
        "відповідь агента має бути короткою, підтримувальною та витриманою в репетиторському стилі;",
        "Відповідь агента має бути короткою, підтримувальною та витриманою в репетиторському стилі.",
    )
    replace_paragraph_text(
        doc,
        "повний розв'язок не подається на початковому етапі взаємодії;",
        "Повний розв'язок не подається на початковому етапі взаємодії.",
    )
    replace_paragraph_text(
        doc,
        "після повторних невдалих спроб допускається конкретніша, але не повна підказка;",
        "Після повторних невдалих спроб допускається конкретніша, але не повна підказка.",
    )
    replace_paragraph_text(
        doc,
        "структура JSON для вправи має містити сигнатуру функції, набір тестів і рубрику оцінювання;",
        "Структура JSON для вправи має містити сигнатуру функції, набір тестів і критерії оцінювання.",
    )
    replace_paragraph_text(
        doc,
        "якщо модель недоступна або повертає невалідний JSON, система переходить у резервний режим.",
        "Якщо модель недоступна або повертає невалідний JSON, система переходить у резервний режим.",
    )
    replace_paragraph_text(
        doc,
        "У роботі виконано функціональну перевірку прототипу, що охоплювала модульні тести, контрольний сценарій перевірки, браузерну функціональну перевірку та локальний виклик Ollama з моделлю Llama 3.1:8b. За її результатами підтверджено коректну ескалацію підказок, явний режим пояснення, генерацію вправи локальною моделлю, полегшений JSON-експорт сесії та стійкість прототипу до невалідних LLM-відповідей. Освітні результати, UX-оцінки та ефективність навчання в межах курсової роботи не оцінювалися.",
        "У роботі виконано функціональну перевірку прототипу, що охоплювала модульні тести, сценарну функціональну перевірку, браузерну функціональну перевірку та перевірку роботи локальної моделі Ollama на базі Llama 3.1:8b. За її результатами підтверджено коректну ескалацію підказок, явний режим пояснення, генерацію вправи за допомогою локальної моделі, експорт даних сесії у форматі JSON та стійкість прототипу до невалідних відповідей LLM. Освітні результати, UX-оцінки та ефективність навчання в межах курсової роботи не оцінювалися.",
    )
    replace_paragraph_text(
        doc,
        "Отже, інженерна функціональність мінімального практичного сценарію підтверджена як модульними тестами, так і повним браузерним сценарієм із локальною Ollama. Під час цього прогону було виявлено й усунуто два локальні недоліки: збільшено тайм-аут LLM-виклику для llama3.1:8b і виправлено скидання firstFailure під час переходу до нової вправи.",
        "Отже, працездатність мінімального практичного сценарію підтверджено як модульними тестами, так і повним браузерним сценарієм з використанням локальної моделі Ollama. У процесі технічної перевірки було виявлено й усунуто два локальні недоліки: збільшено тайм-аут звернення до моделі llama3.1:8b і виправлено скидання ознаки першої зафіксованої помилки під час переходу до нової вправи.",
    )
    replace_paragraph_text(
        doc,
        "На момент завершення роботи наявні лише результати технічної перевірки: модульні тести, контрольний сценарій перевірки, браузерна функціональна перевірка та локальний виклик Ollama з моделлю Llama 3.1:8b.",
        "На момент завершення роботи наявні лише результати технічної перевірки: модульні тести, сценарна функціональна перевірка, браузерна функціональна перевірка та перевірка роботи локальної моделі Ollama на базі Llama 3.1:8b.",
    )
    replace_paragraph_text(
        doc,
        "Цих результатів перевірки достатньо, щоб стверджувати: у репозиторії реалізовано мінімальний практичний сценарій, який структурно відповідає описаній архітектурі та відпрацьовує повний цикл взаємодії з локальною моделлю. Зокрема підтверджено генерацію вправи, ескалацію підказок після повторної помилки, явний режим пояснення та коректний JSON-експорт сесії.",
        "Цих результатів перевірки достатньо, щоб стверджувати: у репозиторії реалізовано мінімальний практичний сценарій, який структурно відповідає описаній архітектурі та відпрацьовує повний цикл взаємодії з локальною моделлю. Зокрема підтверджено генерацію вправи, ескалацію підказок після повторної помилки, явний режим пояснення та коректний експорт даних сесії у форматі JSON.",
    )
    replace_paragraph_text(
        doc,
        "проаналізувати класичні підходи до побудови ITS/АІНС та виділити їхні обов'язкові архітектурні компоненти;",
        "Проаналізувати класичні підходи до побудови ITS/АІНС та виділити їхні обов'язкові архітектурні компоненти;",
    )
    replace_paragraph_text(
        doc,
        "узагальнити можливості й обмеження великих мовних моделей у навчальних сценаріях;",
        "Узагальнити можливості й обмеження великих мовних моделей у навчальних сценаріях;",
    )
    replace_paragraph_text(
        doc,
        "сформулювати функціональні та нефункціональні вимоги до системи;",
        "Сформулювати функціональні та нефункціональні вимоги до системи;",
    )
    replace_paragraph_text(
        doc,
        "запроєктувати модель студента, модель знань, стратегію репетитора, потоки даних і організацію пам'яті агента;",
        "Запроєктувати модель студента, модель знань, стратегію репетитора, потоки даних і організацію пам'яті агента;",
    )
    replace_paragraph_text(
        doc,
        "реалізувати мінімальний браузерний прототип для курсового практичного сценарію з локальною LLM та безпечним тестовим запуском коду;",
        "Реалізувати мінімальний браузерний прототип для курсового практичного сценарію з локальною LLM та безпечним тестовим запуском коду;",
    )
    replace_paragraph_text(
        doc,
        "запропонувати проєкт експериментальної перевірки для подальшого педагогічного оцінювання.",
        "Запропонувати проєкт експериментальної перевірки для подальшого педагогічного оцінювання.",
    )

    assert doc.tables[8].cell(0, 0).text == "Шар"
    assert doc.tables[9].cell(0, 0).text == "Модуль"
    assert doc.tables[10].cell(0, 0).text == "Об'єкт перевірки"

    set_cell_text(
        doc.tables[8].cell(3, 2),
        "простий і прозорий варіант для мінімального практичного сценарію та майбутнього експорту даних сесії у форматі JSON",
    )
    set_cell_text(
        doc.tables[9].cell(2, 2),
        "Зберігає стан користувача, поточну вправу, історію повідомлень, першу зафіксовану помилку та формує експорт даних сесії у форматі JSON.",
    )
    set_cell_text(
        doc.tables[10].cell(2, 1),
        "модульні тести, сценарна функціональна перевірка, браузерна функціональна перевірка",
    )
    set_cell_text(
        doc.tables[10].cell(5, 1),
        "модульні тести, сценарна функціональна перевірка, браузерна функціональна перевірка",
    )
    set_cell_text(doc.tables[10].cell(6, 0), "Швидка сценарна перевірка")
    set_cell_text(doc.tables[10].cell(6, 1), "сценарна функціональна перевірка")
    set_cell_text(
        doc.tables[10].cell(6, 3),
        "Сценарій без важкого стеку швидко перевіряє потік пояснення та формування експорту даних сесії у форматі JSON.",
    )
    set_cell_text(doc.tables[10].cell(7, 0), "Повний сценарій з локальною моделлю Ollama")
    set_cell_text(
        doc.tables[10].cell(7, 1),
        "браузерна функціональна перевірка з використанням локальної моделі Ollama",
    )
    set_cell_text(
        doc.tables[10].cell(7, 3),
        "Підтверджено: генерацію вправи за допомогою локальної моделі Ollama, дві послідовні невдалі спроби з ескалацією підказки, явний режим пояснення та завантаження експорту даних сесії у форматі JSON.",
    )


def export_pngs(pdf_path: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(pdf_path)
    for page_number in range(pdf.page_count):
        pix = pdf[page_number].get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
        pix.save(str(output_dir / f"page_{page_number + 1:02}.png"))
    pdf.close()


def create_preview_sheet(output_path: Path, page_image_paths: list[Path], cols: int = 2) -> None:
    images = [Image.open(path).convert("RGB") for path in page_image_paths]
    thumb_width = 560
    rendered = []
    for image_path, image in zip(page_image_paths, images):
        scale = thumb_width / image.width
        thumb = image.resize((thumb_width, int(image.height * scale)))
        canvas = Image.new("RGB", (thumb.width, thumb.height + 44), "white")
        canvas.paste(thumb, (0, 44))
        draw = ImageDraw.Draw(canvas)
        draw.text((12, 12), image_path.stem.replace("_", " ").title(), fill="black")
        rendered.append(ImageOps.expand(canvas, border=2, fill="black"))

    rows = (len(rendered) + cols - 1) // cols
    cell_width = max(image.width for image in rendered)
    cell_height = max(image.height for image in rendered)
    sheet = Image.new(
        "RGB",
        (40 + cols * cell_width + (cols - 1) * 20, 40 + rows * cell_height + (rows - 1) * 20),
        "#dddddd",
    )

    for index, image in enumerate(rendered):
        row = index // cols
        col = index % cols
        x = 20 + col * (cell_width + 20)
        y = 20 + row * (cell_height + 20)
        sheet.paste(image, (x, y))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(output_path, quality=85)


def create_previews(png_dir: Path, preview_dir: Path) -> None:
    preview_dir.mkdir(parents=True, exist_ok=True)
    groups = {
        "sheet_front.jpg": [1, 2, 3, 4],
        "sheet_concept.jpg": [18, 19, 20, 21],
        "sheet_experiment.jpg": [28, 29, 30, 31, 32, 33],
        "sheet_appendix.jpg": [39, 40, 41, 42, 43, 44],
    }
    for filename, pages in groups.items():
        page_paths = [png_dir / f"page_{page:02}.png" for page in pages]
        create_preview_sheet(preview_dir / filename, page_paths)


def scan_forbidden_terms(doc: Document) -> dict[str, list[str]]:
    hits: dict[str, list[str]] = {}
    for term in FORBIDDEN_TERMS:
        term_hits = []
        for paragraph in doc.paragraphs:
            if term.lower() in paragraph.text.lower():
                term_hits.append(normalize(paragraph.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if term.lower() in paragraph.text.lower():
                            term_hits.append(normalize(paragraph.text))
        if term_hits:
            hits[term] = term_hits
    return hits


def main() -> None:
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(SOURCE_DOC)

    RENDER_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(SOURCE_DOC)
    apply_content_edits(doc)
    doc.save(TARGET_DOC)

    changed = True
    metrics: dict[str, int] = {}
    page_map: dict[str, int] = {}
    for _ in range(3):
        render_pdf(TARGET_DOC, PDF_PATH)
        metrics, page_map, changed = sync_counts_and_toc(TARGET_DOC, PDF_PATH)
        if not changed:
            break

    render_pdf(TARGET_DOC, PDF_PATH)
    final_doc = Document(TARGET_DOC)
    forbidden_hits = scan_forbidden_terms(final_doc)
    export_pngs(PDF_PATH, PNG_DIR)
    create_previews(PNG_DIR, PREVIEW_DIR)

    key_titles = [
        "СПИСОК УМОВНИХ ПОЗНАЧЕНЬ",
        "ВСТУП",
        "3.4. Модель знань і навчального контенту",
        "3.5.1. Формальне визначення агента-репетитора",
        "3.5.2. Архітектура агента - компонентна модель",
        "3.5.3. Матриця рішень і стратегія мінімальної підказки",
        "4.4.1. Функціональна перевірка прототипу",
        "4.4.2. Дизайн педагогічного експерименту",
        "4.5. Аналіз результатів",
        "ВИСНОВКИ",
        "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ",
        "ДОДАТКИ",
        "Додаток А. Діаграми архітектури системи",
        "Додаток Б. Приклади промптів",
        "Додаток В. Приклади діалогів агента",
        "Додаток Г. Фрагменти програмного коду",
    ]

    print(f"TARGET_DOC={TARGET_DOC}")
    print(f"PDF_PATH={PDF_PATH}")
    print(f"PNG_DIR={PNG_DIR}")
    print(f"PREVIEW_DIR={PREVIEW_DIR}")
    print(f"PAGES={metrics['pages']}")
    print(f"FIGURES={metrics['figures']}")
    print(f"TABLES={metrics['tables']}")
    print(f"APPENDICES={metrics['appendices']}")
    print(f"REFERENCES={metrics['references']}")
    for title in key_titles:
        print(f"PAGE::{title}={page_map[title]}")
    if forbidden_hits:
        print("FORBIDDEN_TERMS_FOUND")
        for term, hits in forbidden_hits.items():
            print(term, "=>", hits[:5])
    else:
        print("FORBIDDEN_TERMS_FOUND=0")


if __name__ == "__main__":
    main()
