from __future__ import annotations

import subprocess
from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
INPUT_PATH = ROOT / "output" / "coursework_draft_ua.docx"
OUTPUT_PATH = ROOT / "output" / "coursework_draft_ua_fixed.docx"
WORK_DIR = ROOT / "tmp" / "docs" / "fixed"
DIAGRAM_DIR = ROOT / "docs" / "codex" / "diagram_assets"
TEMP_STAGE_ONE = WORK_DIR / "coursework_stage1.docx"
TEMP_STAGE_TWO = WORK_DIR / "coursework_stage2.docx"
TEMP_PDF_ONE = WORK_DIR / "coursework_stage1.pdf"
TEMP_PDF_TWO = WORK_DIR / "coursework_stage2.pdf"
FINAL_PDF = WORK_DIR / "coursework_draft_ua_fixed.pdf"

ACCESS_DATE = "14.03.2026"
FONT_REGULAR = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
FONT_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")

TABLE_CAPTIONS = [
    "Таблиця 1.1 - Порівняння сучасних систем-репетиторів на основі ШІ.",
    "Таблиця 3.1 - Функціональні вимоги до прототипу АІНС.",
    "Таблиця 3.2 - Нефункціональні вимоги до прототипу АІНС.",
    "Таблиця 3.3 - Компоненти запропонованої архітектури АІНС.",
    "Таблиця 3.4 - Поля спрощеної моделі стану студента `UserState`.",
    "Таблиця 3.5 - Матриця педагогічних дій агента-репетитора.",
    "Таблиця 3.6 - Організація пам'яті агента.",
    "Таблиця 4.1 - Технологічний стек прототипу.",
    "Таблиця 4.2 - Реалізовані модулі прототипу.",
    "Таблиця 4.3 - Результати функціональної валідації прототипу.",
    "Таблиця 4.4 - Гіпотези педагогічного експерименту.",
    "Таблиця 4.5 - План параметрів та метрик педагогічного експерименту.",
    "Таблиця 4.6 - Операціоналізація змінних педагогічного експерименту.",
    "Таблиця 4.7 - Шаблон таблиці для ручного занесення результатів експерименту.",
]

REPLACEMENTS = {
    "{cite('S11', 'S12')}": "[11], [12]",
    "browser-based": "браузерний",
    "fallback-логікою": "резервною логікою",
    "fallback-логіка": "резервна логіка",
    "fallback-вправи": "резервні вправи",
    "fallback-набір": "резервний набір",
    "fallback-режим": "резервний режим",
    "fallback-механізм": "резервний механізм",
    "figure placeholders": "текстові місця для рисунків",
    "prompt engineering": "проєктування промптів",
    "post-test": "підсумковий тест",
    "teacher assistant": "асистент викладача",
    "coding tasks": "задач для програмування",
}

FIGURE_CAPTIONS = {
    "3.1": "Рисунок 3.1 - Компонентна архітектура АІНС для сценарію навчання базового JavaScript.",
    "3.2": "Рисунок 3.2 - Структура `UserState` та зв'язок із `SessionMemory`.",
    "3.3": "Рисунок 3.3 - Діаграма послідовності для сценарію «згенерувати вправу -> пройти тести -> дати фідбек».",
    "4.1": "Рисунок 4.1 - Схема виконання прототипу з локальною LLM, резервним режимом і `Web Worker`.",
}

SUPPLEMENTAL_KEY_REPLACEMENTS = {
    "У роботі систематизовано класичні компоненти ITS, проаналізовано сучасні підходи до використання LLM в освіті, запропоновано архітектуру агента-репетитора з моделлю стану студента, матрицею рішень і багаторівневою організацією пам'яті, а також реалізовано мінімальний браузерний прототип із локальним LLM runtime, Web Worker sandbox та резервною логікою.": "У роботі систематизовано класичні компоненти ITS, проаналізовано сучасні підходи до використання LLM в освіті, запропоновано архітектуру агента-репетитора з моделлю стану студента, матрицею рішень і багаторівневою організацією пам'яті, а також реалізовано мінімальний браузерний прототип із локальною LLM, `Web Worker` для тестів та резервною логікою.",
    "реалізувати мінімальний браузерний прототип для курсового vertical slice з локальним LLM runtime та безпечним тестовим запуском коду;": "реалізувати мінімальний браузерний прототип для курсового практичного сценарію з локальною LLM та безпечним тестовим запуском коду;",
    "У центрі персоналізації перебуває модель студента. Повноцінні ITS використовують для цього складніші probabilistic або knowledge tracing механізми, але для vertical slice доцільно застосувати спрощений вектор стану `UserState`, який усе ж зберігає ключові сигнали для педагогічного рішення [2], [3], [4].": "У центрі персоналізації перебуває модель студента. Повноцінні ITS використовують для цього складніші probabilistic або knowledge tracing механізми, але для мінімального практичного сценарію доцільно застосувати спрощений вектор стану `UserState`, який усе ж зберігає ключові сигнали для педагогічного рішення [2], [3], [4].",
    "Оскільки в репозиторії на початку роботи не було прикладної реалізації, практичну частину свідомо обмежено до одного мінімального vertical slice. Метою було не створити повну платформу, а отримати реальний кодовий артефакт, що демонструє архітектурні рішення з розділу 3.": "Оскільки в репозиторії на початку роботи не було прикладної реалізації, практичну частину свідомо обмежено до одного мінімального практичного сценарію. Метою було не створити повну платформу, а отримати реальний кодовий артефакт, що демонструє архітектурні рішення з розділу 3.",
    "Отже, інженерна функціональність vertical slice підтверджена частково: чисті модулі й базова статична доставка працюють, але повний інтерактивний сценарій з локальним Ollama має бути додатково перевірений вручну.": "Отже, інженерна функціональність мінімального практичного сценарію підтверджена частково: чисті модулі й базова статична доставка працюють, але повний інтерактивний сценарій з локальним Ollama має бути додатково перевірений вручну.",
    "На момент завершення автоматично згенерованої чернетки наявні лише результати інженерної валідації. Їх достатньо, щоб стверджувати: у репозиторії присутній реальний vertical slice, який структурно відповідає описаній архітектурі, але вони недостатні для висновків про педагогічну ефективність.": "На момент завершення автоматично згенерованої чернетки наявні лише результати інженерної валідації. Їх достатньо, щоб стверджувати: у репозиторії присутній реальний мінімальний практичний сценарій, який структурно відповідає описаній архітектурі, але вони недостатні для висновків про педагогічну ефективність.",
    "Для задачі базового JavaScript tutor не потрібна максимальна генеративна потужність, натомість важливі локальність, передбачуваність і можливість працювати з невеликим структурованим JSON-контрактом. Саме тому в роботі обрано локальний runtime `Ollama + llama3.1:8b` як цільове рішення для прототипу [25].": "Для задачі базового JavaScript tutor не потрібна максимальна генеративна потужність, натомість важливі локальність, передбачуваність і можливість працювати з невеликим структурованим JSON-контрактом. Саме тому в роботі обрано локальне середовище виконання `Ollama + llama3.1:8b` як цільове рішення для прототипу [25].",
    "Вибір локальної Llama через Ollama зумовлено кількома факторами. По-перше, це відповідає вимозі використовувати безкоштовний локальний runtime без звернення до платних хмарних API. По-друге, Ollama надає достатньо простий HTTP API для інтеграції у lightweight-прототип. По-третє, модель класу 8B достатня для генерації невеликих задач, коротких підказок та структурованих відповідей у межах базового навчального сценарію [17], [18], [19], [25].": "Вибір локальної Llama через Ollama зумовлено кількома факторами. По-перше, це відповідає вимозі використовувати безкоштовне локальне середовище виконання без звернення до платних хмарних API. По-друге, Ollama надає достатньо простий HTTP API для інтеграції у lightweight-прототип. По-третє, модель класу 8B достатня для генерації невеликих задач, коротких підказок та структурованих відповідей у межах базового навчального сценарію [17], [18], [19], [25].",
    "У пропонованій АІНС LLM виконує роль ядра модуля генерації, але не є єдиним носієм логіки системи. Вона перетворює структурований контекст у зрозумілу підказку або вправу, однак сама політика, модель стану та безпечний execution runtime мають бути винесені за межі моделі. Такий підхід дозволяє зменшити вплив галюцинацій і краще контролювати освітню доцільність відповіді [9], [10], [13].": "У пропонованій АІНС LLM виконує роль ядра модуля генерації, але не є єдиним носієм логіки системи. Вона перетворює структурований контекст у зрозумілу підказку або вправу, однак сама політика, модель стану та безпечне середовище виконання мають бути винесені за межі моделі. Такий підхід дозволяє зменшити вплив галюцинацій і краще контролювати освітню доцільність відповіді [9], [10], [13].",
    "У поточному середовищі локальний `ollama` binary не був встановлений, тому прототип додатково отримав деградаційний режим. Це важливий практичний висновок: для дослідницької системи потрібно з самого початку передбачати fallback-логіку, а не вважати зовнішній runtime гарантовано доступним [25].": "У поточному середовищі локальний `ollama` binary не був встановлений, тому прототип додатково отримав резервний режим. Це важливий практичний висновок: для дослідницької системи потрібно з самого початку передбачати резервну логіку, а не вважати зовнішнє середовище виконання гарантовано доступним [25].",
    "Ключове обмеження полягає в тому, що локальний Ollama runtime не був встановлений у поточному середовищі. Це означає, що частина практичної логіки перевірялася у резервному режимі, а повна робота з реальною локальною моделлю потребує ручного запуску.": "Ключове обмеження полягає в тому, що локальне середовище виконання Ollama не було встановлене у поточному середовищі. Це означає, що частина практичної логіки перевірялася у резервному режимі, а повна робота з реальною локальною моделлю потребує ручного запуску.",
    "Для практичної частини реалізовано мінімальний vertical slice JavaScript tutor: генерація вправи, введення коду, запуск тестів у Web Worker, адаптивний фідбек, SessionMemory та резервна логіка у випадку недоступності або нестабільності локальної LLM.": "Для практичної частини реалізовано мінімальний практичний сценарій JavaScript tutor: генерація вправи, введення коду, запуск тестів у `Web Worker`, адаптивний фідбек, `SessionMemory` та резервна логіка у випадку недоступності або нестабільності локальної LLM.",
    "3.5.3. Матриця рішень і minimal feedback strategy": "3.5.3. Матриця рішень і стратегія мінімальної підказки",
    "4.3. Розробка prompt-стратегій": "4.3. Проєктування промптів для генерації вправ і фідбеку",
}


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(str(font_path), size=size)


def mm_to_emu(value_mm: float) -> int:
    return int(value_mm * 36000)


def normalize_spaces(text: str) -> str:
    return " ".join(text.replace("\n", " ").split())


def set_run_font(run, size=14, bold=False, italic=False, font_name="Times New Roman") -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def apply_font_to_paragraph(paragraph: Paragraph, size=14, bold=False, italic=False, align=None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)
    if align is not None:
        paragraph.alignment = align


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("pPr"):
            continue
        paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, size=14, bold=False, italic=False, align=None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    if align is not None:
        paragraph.alignment = align


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def iter_block_items(document: Document):
    for child in document.element.body:
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def normalize_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(215.9)
    section.page_height = Mm(279.4)
    section.left_margin = Mm(20)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(14)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.first_line_indent = Mm(12.5)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Title", "List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = style_name.startswith("Heading") or style_name == "Title"
        style.font.color.rgb = None
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for paragraph in document.paragraphs:
        if paragraph.style.name == "Normal":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif paragraph.style.name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif paragraph.style.name in {"Heading 2", "Heading 3"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_font_to_paragraph(
            paragraph,
            size=14,
            bold=paragraph.style.name.startswith("Heading") or paragraph.style.name == "Title",
        )

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    apply_font_to_paragraph(
                        paragraph,
                        size=12,
                        bold=row_index == 0,
                    )


def contains_cyrillic(text: str) -> bool:
    return any("А" <= char <= "я" or char in "ЇїІіЄєҐґ" for char in text)


def apply_replacements(text: str) -> str:
    if "{cite(" not in text and "docs/codex/" not in text and not contains_cyrillic(text):
        return text
    updated = text
    for source, target in REPLACEMENTS.items():
        updated = updated.replace(source, target)
    return updated


def update_textual_content(document: Document) -> None:
    for paragraph in document.paragraphs:
        if not paragraph.text:
            continue
        updated = apply_replacements(paragraph.text)
        if updated != paragraph.text:
            set_paragraph_text(
                paragraph,
                updated,
                size=10 if "Courier New" in paragraph._p.xml else 14,
                bold=paragraph.style.name.startswith("Heading"),
                align=paragraph.alignment,
            )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    updated = apply_replacements(paragraph.text)
                    if updated != paragraph.text:
                        set_paragraph_text(
                            paragraph,
                            updated,
                            size=12,
                            bold=row == table.rows[0],
                            align=paragraph.alignment,
                        )


def add_page_number_field(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=12)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def configure_page_numbers(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.different_first_page_header_footer = True
    add_page_number_field(section.header.paragraphs[0])
    clear_paragraph(section.first_page_header.paragraphs[0])
    clear_paragraph(section.footer.paragraphs[0])
    clear_paragraph(section.first_page_footer.paragraphs[0])


def normalize_section_breaks(document: Document) -> None:
    blocks = list(iter_block_items(document))
    for index, block in enumerate(blocks):
        if not isinstance(block, Paragraph):
            continue
        if block.style.name != "Heading 1":
            continue
        if index == 0:
            continue
        block.paragraph_format.page_break_before = True
        previous = blocks[index - 1]
        if (
            isinstance(previous, Paragraph)
            and not previous.text.strip()
            and ("w:br" in previous._p.xml or "lastRenderedPageBreak" in previous._p.xml)
        ):
            remove_paragraph(previous)


def insert_paragraph_before_table(table: Table, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addprevious(new_p)
    paragraph = Paragraph(new_p, table._parent)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Mm(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)
    return paragraph


def add_table_captions(document: Document) -> None:
    for table, caption in zip(document.tables, TABLE_CAPTIONS):
        previous = table._tbl.getprevious()
        if previous is not None and previous.tag.endswith("}p"):
            paragraph = Paragraph(previous, table._parent)
            if paragraph.text.strip().startswith("Таблиця "):
                set_paragraph_text(paragraph, caption, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
                continue
        insert_paragraph_before_table(table, caption)


def populate_validation_table(table: Table) -> None:
    rows = [
        ["Політика підказок", "`npm test`", "пройдено", "Перевірено: перша помилка -> `minimal_hint`; повторна та сама помилка -> `targeted_hint`."],
        ["Fallback explanation", "`npm test`", "пройдено", "Перевірено, що резервне пояснення не дублює студентський запит."],
        ["Schema validation JSON", "`npm test`", "пройдено", "Перевірено `null`, масив, рядок і частковий об'єкт для exercise/feedback JSON."],
        ["Sandbox worker", "`npm test`", "пройдено", "Перевірено `missing function`, timeout path, `functionName` validation, `worker` error handling."],
        ["Синтаксис JS-модулів", "`node --check` по `prototype/src` і `prototype/tests`", "пройдено", "Помилок парсингу не виявлено."],
        ["Static smoke check", "`npm run serve` + `curl -I /prototype/`", "manual pending", "Може бути виконано локально без важкого стеку; в автоматичному циклі не запускався окремо."],
        ["E2E з локальною Ollama", "ручний запуск у браузері", "manual pending", "Потребує встановленого `ollama` і локальної моделі `llama3.1:8b`."],
    ]

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value


def populate_module_table(table: Table) -> None:
    rows = [
        ["Chat UI", "`prototype/index.html`, `prototype/src/main.js`", "Інтерфейс вибору теми, редактор коду, журнал взаємодій, синхронізація моделі та дружні повідомлення про помилки."],
        ["SessionMemory", "`prototype/src/state/sessionMemory.js`", "Зберігає `UserState`, поточну вправу, історію повідомлень та сигнатури помилок."],
        ["TutorPolicy", "`prototype/src/services/tutorPolicy.js`", "Rule-based політика з коректною ескалацією: перша помилка -> `minimal_hint`, повторна та сама -> `targeted_hint`."],
        ["ExerciseGenerator", "`prototype/src/services/exerciseGenerator.js`", "Генерація вправи через LLM, явна schema validation і перехід до резервної вправи при невідповідному JSON."],
        ["ClientTestRunner", "`prototype/src/services/clientTestRunner.js`, `prototype/src/workers/jsSandboxWorker.js`", "Клієнтський запуск тестів у `Web Worker` з timeout, validation payload, `worker.onerror` і `messageerror`."],
        ["FeedbackEvaluator", "`prototype/src/services/feedbackEvaluator.js`", "Адаптивний фідбек через LLM або резервний режим; fallback explanation формує змістовне пояснення концепту."],
        ["OllamaClient", "`prototype/src/services/ollamaClient.js`", "HTTP API, перевірка доступності, schema validation, retry/fallback для JSON-відповідей."],
    ]

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value


def update_core_tables(document: Document) -> None:
    populate_module_table(document.tables[8])
    populate_validation_table(document.tables[9])

    hypothesis_table = document.tables[10]
    while len(hypothesis_table.rows) > 1:
        hypothesis_table._tbl.remove(hypothesis_table.rows[-1]._tr)
    for hypothesis, description in [
        ["H1", "Адаптивний агент покращує результат підсумкового тесту порівняно з режимом без адаптивного фідбеку."],
        ["H2", "Адаптивний агент скорочує кількість безрезультатних спроб і час до коректного розв'язку."],
        ["H3", "Адаптивний агент підвищує суб'єктивне відчуття підтримки без зростання залежності від готових відповідей."],
    ]:
        cells = hypothesis_table.add_row().cells
        cells[0].text = hypothesis
        cells[1].text = description

    plan_table = document.tables[11]
    while len(plan_table.rows) > 1:
        plan_table._tbl.remove(plan_table.rows[-1]._tr)
    for parameter, plan in [
        ["Дизайн", "crossover або A/B із попереднім і підсумковим тестом; обидві умови проходять одну й ту саму тематичну міні-сесію."],
        ["Цільова вибірка", "24-40 студентів, які вивчають базові теми JavaScript."],
        ["Основні метрики", "правильність розв'язку, час виконання, кількість спроб, якість самостійного пояснення."],
        ["Додаткові метрики", "самооцінка впевненості, суб'єктивна корисність підказок, потреба в повторному поясненні."],
        ["Статистичний аналіз", "paired t-test або Wilcoxon; effect size розраховувати лише після отримання реальних даних."],
    ]:
        cells = plan_table.add_row().cells
        cells[0].text = parameter
        cells[1].text = plan


def build_table(document: Document, header: list[str], rows: list[list[str]]) -> Table:
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.text = value

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                apply_font_to_paragraph(paragraph, size=12, bold=row_index == 0)

    return table


def insert_table_before_paragraph(
    document: Document,
    anchor: Paragraph,
    caption: str,
    header: list[str],
    rows: list[list[str]],
) -> None:
    caption_paragraph = anchor.insert_paragraph_before()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Mm(0)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=12)

    table = build_table(document, header, rows)
    anchor._p.addprevious(table._tbl)


def strengthen_experiment_section(document: Document) -> None:
    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ: після проведення реального експерименту")
    )

    statistical_plan = anchor.insert_paragraph_before()
    statistical_plan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = statistical_plan.add_run(
        "Для уникнення спекулятивних висновків у роботі фіксується лише план статистичного аналізу: перевірка нормальності розподілу, вибір параметричного або непараметричного критерію, оцінка величини ефекту та окремий аналіз перенесення знань на delayed test. Числові результати, effect size та p-values мають бути внесені лише після збору реальних даних."
    )
    set_run_font(run)

    insert_table_before_paragraph(
        document,
        statistical_plan,
        "Таблиця 4.6 - Операціоналізація змінних педагогічного експерименту.",
        ["Тип змінної", "Позначення", "Опис"],
        [
            ["Незалежна", "X", "Режим навчання: адаптивний агент або контрольний сценарій без адаптивного фідбеку."],
            ["Залежна", "Y1", "Результат підсумкового тесту з базового JavaScript."],
            ["Залежна", "Y2", "Час до коректного розв'язку."],
            ["Залежна", "Y3", "Кількість спроб до першого коректного розв'язку."],
            ["Контрольована", "C1", "Однакова тема, складність вправи, набір тестів і часовий ліміт сесії."],
        ],
    )

    insert_table_before_paragraph(
        document,
        statistical_plan,
        "Таблиця 4.7 - Шаблон таблиці для ручного занесення результатів експерименту.",
        ["Метрика", "Умова A", "Умова B", "Примітка"],
        [
            ["Середній pre-test", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", ""],
            ["Середній post-test", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", ""],
            ["Середній час до розв'язку", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", ""],
            ["Середня кількість спроб", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", ""],
            ["Статистичний висновок", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", "[ПОТРІБНО РУЧНЕ ВТРУЧАННЯ]", "Заповнюється після розрахунків."],
        ],
    )


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines = []
    current = []

    for word in words:
        candidate = " ".join([*current, word])
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width:
            current.append(word)
        else:
            lines.append(" ".join(current))
            current = [word]

    if current:
        lines.append(" ".join(current))

    return lines


def draw_box(draw, box, text, fill, outline, title_font, body_font, radius=18):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    lines = wrap_text(draw, text, body_font, x2 - x1 - 28)
    total_height = len(lines) * 28
    y = y1 + (y2 - y1 - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=body_font)
        width = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, fill="#1f2933", font=body_font)
        y += 28


def draw_arrow(draw, start, end, fill="#385170", width=5):
    draw.line([start, end], fill=fill, width=width)
    angle_x = end[0] - start[0]
    angle_y = end[1] - start[1]
    if angle_x == 0 and angle_y == 0:
        return
    length = (angle_x ** 2 + angle_y ** 2) ** 0.5
    ux, uy = angle_x / length, angle_y / length
    left = (end[0] - 18 * ux + 10 * uy, end[1] - 18 * uy - 10 * ux)
    right = (end[0] - 18 * ux - 10 * uy, end[1] - 18 * uy + 10 * ux)
    draw.polygon([end, left, right], fill=fill)


def create_canvas(size=(1600, 1000), title="") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, "#ffffff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    subtitle_font = load_font(22, bold=False)
    if title:
        draw.text((60, 32), title, fill="#102a43", font=title_font)
        draw.text((60, 76), "Магістерська курсова робота: архітектура АІНС і вертикальний сценарій для JavaScript", fill="#486581", font=subtitle_font)
    return image, draw


def save_component_diagram(path: Path) -> None:
    image, draw = create_canvas(title="Компонентна архітектура АІНС")
    title_font = load_font(26, bold=True)
    body_font = load_font(22)

    boxes = {
        "student": (70, 250, 260, 360, "Студент", "#fef3c7"),
        "ui": (320, 230, 580, 380, "Chat UI / Робоче навчальне середовище", "#dbeafe"),
        "memory": (660, 120, 930, 270, "SessionMemory / UserState", "#dcfce7"),
        "policy": (980, 120, 1260, 270, "TutorPolicy / модуль педагогічного рішення", "#fee2e2"),
        "exercise": (660, 340, 930, 480, "ExerciseGenerator", "#ede9fe"),
        "feedback": (980, 340, 1260, 480, "FeedbackEvaluator", "#ede9fe"),
        "runner": (660, 560, 930, 700, "ClientTestRunner", "#fce7f3"),
        "worker": (980, 560, 1260, 700, "Web Worker sandbox", "#fce7f3"),
        "ollama": (1320, 250, 1530, 420, "Локальна LLM через Ollama", "#ccfbf1"),
        "knowledge": (660, 780, 1260, 920, "Концептуальні модулі повної АІНС: Knowledge Memory, Summary Memory, Vector/Profile Memory", "#f3f4f6"),
    }

    for _, (x1, y1, x2, y2, text, fill) in boxes.items():
        draw_box(draw, (x1, y1, x2, y2), text, fill, "#334e68", title_font, body_font)

    draw_arrow(draw, (260, 305), (320, 305))
    draw_arrow(draw, (580, 260), (660, 195))
    draw_arrow(draw, (580, 305), (980, 195))
    draw_arrow(draw, (580, 340), (660, 410))
    draw_arrow(draw, (580, 350), (980, 410))
    draw_arrow(draw, (580, 365), (660, 630))
    draw_arrow(draw, (930, 410), (1320, 330))
    draw_arrow(draw, (1260, 410), (1320, 350))
    draw_arrow(draw, (930, 630), (980, 630))
    draw_arrow(draw, (1260, 630), (580, 365))
    draw_arrow(draw, (930, 195), (660, 850))
    draw_arrow(draw, (1120, 270), (1120, 780))

    image.save(path)


def save_sequence_diagram(path: Path) -> None:
    image, draw = create_canvas(title="Послідовність сценарію генерації вправи, тестування і фідбеку")
    body_font = load_font(20)
    header_font = load_font(22, bold=True)

    participants = [
        ("Студент", 120),
        ("UI", 320),
        ("ExerciseGenerator", 540),
        ("ClientTestRunner", 820),
        ("TutorPolicy", 1100),
        ("FeedbackEvaluator", 1320),
    ]

    for label, x in participants:
        draw.rounded_rectangle((x - 80, 120, x + 80, 180), radius=16, fill="#eff6ff", outline="#1d4ed8", width=3)
        bbox = draw.textbbox((0, 0), label, font=header_font)
        draw.text((x - (bbox[2] - bbox[0]) / 2, 138), label, fill="#102a43", font=header_font)
        draw.line((x, 180, x, 900), fill="#94a3b8", width=3)

    steps = [
        (230, 320, 540, "1. Обрати тему і згенерувати вправу"),
        (290, 540, 320, "2. JSON prompt / резервний режим"),
        (360, 320, 120, "3. Показати вправу студенту"),
        (440, 120, 320, "4. Ввести код і запустити тести"),
        (510, 320, 820, "5. Передати код і тести в sandbox"),
        (580, 820, 320, "6. Повернути `runResult` або timeout"),
        (650, 320, 1100, "7. Обрати педагогічну дію"),
        (720, 320, 1320, "8. Запросити адаптивний фідбек"),
        (790, 1320, 320, "9. Повернути JSON фідбек або резервне пояснення"),
        (860, 320, 120, "10. Показати підказку і наступний крок"),
    ]

    for y, x1, x2, label in steps:
        draw_arrow(draw, (x1, y), (x2, y))
        bbox = draw.textbbox((0, 0), label, font=body_font)
        text_x = min(x1, x2) + abs(x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.rectangle((text_x - 12, y - 26, text_x + (bbox[2] - bbox[0]) + 12, y + 8), fill="#ffffff")
        draw.text((text_x, y - 24), label, fill="#1f2937", font=body_font)

    image.save(path)


def save_state_diagram(path: Path) -> None:
    image, draw = create_canvas(title="Структура UserState і SessionMemory")
    body_font = load_font(22)
    header_font = load_font(24, bold=True)

    draw_box(draw, (90, 180, 720, 860), "SessionMemory", "#ecfeff", "#0f766e", header_font, header_font)
    draw.rounded_rectangle((140, 250, 670, 610), radius=18, fill="#ffffff", outline="#0f766e", width=3)
    draw.text((170, 270), "UserState", fill="#134e4a", font=header_font)

    fields = [
        "knowledgeLevel",
        "errorHistory",
        "confidence",
        "currentTopic",
        "attemptsCount",
        "lastAction",
    ]
    for index, field in enumerate(fields):
        y = 330 + index * 42
        draw.text((190, y), f"• {field}", fill="#1f2937", font=body_font)

    draw.rounded_rectangle((140, 650, 670, 790), radius=18, fill="#f8fafc", outline="#64748b", width=3)
    draw.text((170, 675), "Додаткові поля сесії", fill="#334155", font=header_font)
    draw.text((190, 725), "exercise, transcript, lastRunResult, lastErrorSignature", fill="#1f2937", font=body_font)

    draw_box(draw, (880, 260, 1490, 430), "Вхідні сигнали: studentRequest, runResult, topic, code", "#fef2f2", "#b91c1c", header_font, body_font)
    draw_box(draw, (880, 520, 1490, 760), "Оновлення стану:\n• attemptsCount += 1\n• deriveConfidence(runResult)\n• buildRunErrorSignature(runResult)\n• history[:8]", "#eef2ff", "#4338ca", header_font, body_font)

    draw_arrow(draw, (720, 350), (880, 350))
    draw_arrow(draw, (880, 620), (720, 620))

    image.save(path)


def save_runtime_diagram(path: Path) -> None:
    image, draw = create_canvas(title="Схема виконання прототипу")
    body_font = load_font(22)
    header_font = load_font(24, bold=True)

    nodes = [
        ((120, 280, 430, 430), "Браузерний UI", "#dbeafe", "#1d4ed8"),
        ((520, 180, 860, 350), "Локальна LLM через Ollama", "#ccfbf1", "#0f766e"),
        ((520, 520, 860, 690), "Web Worker для тестів", "#fae8ff", "#9333ea"),
        ((960, 180, 1450, 350), "Schema validation + retry / резервний режим", "#fef3c7", "#d97706"),
        ((960, 520, 1450, 690), "Повідомлення про помилки без ламання UX", "#fee2e2", "#b91c1c"),
    ]

    for box, text, fill, outline in nodes:
        draw_box(draw, box, text, fill, outline, header_font, body_font)

    draw_arrow(draw, (430, 330), (520, 260))
    draw_arrow(draw, (430, 370), (520, 600))
    draw_arrow(draw, (860, 260), (960, 260))
    draw_arrow(draw, (860, 600), (960, 600))
    draw_arrow(draw, (1450, 260), (430, 330))
    draw_arrow(draw, (1450, 600), (430, 370))

    legend = dedent(
        """\
        Основні технічні рішення:
        • модель синхронізується перед кожним LLM-викликом;
        • некоректний JSON не валить застосунок, а переводить його в резервний режим;
        • worker має timeout, validation payload і `worker.onerror` / `messageerror`.
        """
    ).strip()
    draw.multiline_text((120, 770), legend, fill="#1f2937", font=body_font, spacing=10)

    image.save(path)


def generate_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "3.1": DIAGRAM_DIR / "figure_3_1_component_architecture.png",
        "3.2": DIAGRAM_DIR / "figure_3_2_session_memory.png",
        "3.3": DIAGRAM_DIR / "figure_3_3_sequence.png",
        "4.1": DIAGRAM_DIR / "figure_4_1_runtime.png",
    }
    save_component_diagram(paths["3.1"])
    save_state_diagram(paths["3.2"])
    save_sequence_diagram(paths["3.3"])
    save_runtime_diagram(paths["4.1"])
    return paths


def insert_picture_before(paragraph: Paragraph, image_path: Path, width_mm: float = 155) -> None:
    picture_paragraph = paragraph.insert_paragraph_before()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.first_line_indent = Mm(0)
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))


def attach_figures(document: Document, diagram_paths: dict[str, Path]) -> None:
    paragraph_map = {normalize_spaces(p.text): p for p in document.paragraphs if p.text.strip()}

    replace_map = {
        "Рисунок 3.1 - Загальна компонентна архітектура АІНС для JavaScript tutor.": ("3.1", 154),
        "Рисунок 3.2 - Структура UserState та її зв'язок із SessionMemory.": ("3.2", 150),
        "Рисунок 3.3 - Цикл роботи агента за патерном Reason -> Act -> Observe.": ("3.3", 156),
    }

    for old_text, (figure_key, width_mm) in replace_map.items():
        paragraph = paragraph_map.get(normalize_spaces(old_text))
        if paragraph is None:
            continue
        previous = paragraph._p.getprevious()
        if previous is not None and previous.tag.endswith("}p"):
            previous_paragraph = Paragraph(previous, paragraph._parent)
            if "Плейсхолдер для рисунка" in previous_paragraph.text:
                remove_paragraph(previous_paragraph)
        set_paragraph_text(paragraph, FIGURE_CAPTIONS[figure_key], size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        insert_picture_before(paragraph, diagram_paths[figure_key], width_mm=width_mm)

    section_4_anchor = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "4.2. Реалізація ключових модулів системи"
    )
    figure_intro = section_4_anchor.insert_paragraph_before()
    figure_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = figure_intro.add_run(
        "На рисунку 4.1 показано, як у прототипі розведено браузерний інтерфейс, локальний LLM-виклик, резервний режим і клієнтський запуск тестів у `Web Worker`."
    )
    set_run_font(run)
    runtime_caption = section_4_anchor.insert_paragraph_before()
    set_paragraph_text(runtime_caption, FIGURE_CAPTIONS["4.1"], size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    insert_picture_before(runtime_caption, diagram_paths["4.1"], width_mm=156)


def collect_headings(document: Document) -> list[dict[str, str | int]]:
    entries = []
    for paragraph in document.paragraphs:
        if paragraph.style.name not in {"Heading 1", "Heading 2", "Heading 3"}:
            continue
        text = paragraph.text.strip()
        if not text or text == "ЗМІСТ":
            continue
        entries.append(
            {
                "text": text,
                "level": {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}[paragraph.style.name],
            }
        )
    return entries


def replace_toc(document: Document, headings: list[dict[str, str | int]], page_map: dict[str, int] | None) -> None:
    toc_heading = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "ЗМІСТ")
    next_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "СПИСОК УМОВНИХ ПОЗНАЧЕНЬ"
    )

    current = toc_heading._p.getnext()
    while current is not None and current is not next_heading._p:
        next_current = current.getnext()
        toc_heading._p.getparent().remove(current)
        current = next_current

    for entry in headings:
        paragraph = next_heading.insert_paragraph_before()
        paragraph.paragraph_format.first_line_indent = Mm(0)
        paragraph.paragraph_format.left_indent = Mm(0 if entry["level"] == 1 else 8 if entry["level"] == 2 else 16)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Mm(175), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        label = entry["text"] if page_map else "..." * max(3, int(len(entry["text"]) / 3))
        page_number = str(page_map.get(entry["text"], "")) if page_map else "00"
        run = paragraph.add_run(f"{label}\t{page_number}")
        set_run_font(run)


def render_docx_to_pdf(input_path: Path, output_pdf: Path) -> None:
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            "-env:UserInstallation=file:///tmp/lo_profile_coursework_fix",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_pdf.parent),
            str(input_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )


def extract_page_map(pdf_path: Path, headings: list[dict[str, str | int]]) -> dict[str, int]:
    reader = PdfReader(str(pdf_path))
    normalized_pages = [normalize_spaces(page.extract_text() or "") for page in reader.pages]
    page_map = {}
    for entry in headings:
        heading_text = normalize_spaces(entry["text"])
        for page_number, page_text in enumerate(normalized_pages, start=1):
            if heading_text in page_text:
                page_map[entry["text"]] = page_number
                break
    return page_map


def update_annotation_counts(document: Document, page_count: int) -> None:
    figure_count = len(document.inline_shapes)
    table_count = len(document.tables)
    appendix_count = 4
    source_count = 29

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Магістерська курсова робота:"):
            set_paragraph_text(
                paragraph,
                f"Магістерська курсова робота: {page_count} с., {figure_count} рис., {table_count} табл., {appendix_count} дод., {source_count} джерел.",
            )
        if text.startswith("Master's coursework project:"):
            set_paragraph_text(
                paragraph,
                f"Master's coursework project: {page_count} pages, {figure_count} figures, {table_count} tables, {appendix_count} appendices, and {source_count} references.",
            )


def update_key_paragraphs(document: Document) -> None:
    replacements = {
        "Текстові джерела для діаграм винесено в окремий файл `docs/codex/DIAGRAM_SOURCES.md`, що дозволяє надалі замінити плейсхолдери на відрендерені схеми без зміни структури документа.": "На рисунку 3.1 узагальнено взаємодію концептуальних і реалізованих модулів АІНС. До практичної частини увійшли лише ті компоненти, які реально присутні в репозиторії.",
        "Тут `K` позначає оцінку знань, `E` - історію помилок, `C` - впевненість, `Topic` - поточну тему, `Attempts` - кількість спроб, `M` - пам'ять сесії, `O` - нові спостереження, `A` - множину дій агента, `T` - функцію переходу стану, `P` - політику, а `R` - функцію генерації відповіді. Такий опис дозволяє відокремити міркування агента від суто мовної генерації [11], [12].": "Тут `K` позначає оцінку знань, `E` - історію помилок, `C` - впевненість, `Topic` - поточну тему, `Attempts` - кількість спроб, `M` - пам'ять сесії, `O` - нові спостереження, `A` - множину дій агента, `T` - функцію переходу стану, `P` - політику, а `R` - функцію генерації відповіді. Такий опис дозволяє відокремити логіку міркування агента від мовної генерації [11], [12].",
        "Ключова вимога до агента-репетитора полягає в тому, щоб допомога була мінімально достатньою. Система не повинна відразу видавати готовий код, якщо студент ще може просунутися самостійно. Тому політика будується за принципом ескалації підказок залежно від кількості спроб та повторюваності помилки.": "Ключова вимога до агента-репетитора полягає в тому, щоб допомога була мінімально достатньою. Після технічного аудиту правило ескалації було уточнено: перша невдала спроба дає `minimal_hint`, а `targeted_hint` з'являється лише після повторення тієї самої помилки або серії кількох невдалих спроб.",
        "У поточній чернетці в основних розділах використано текстові місця для рисунків. Текстові джерела діаграм збережено у файлі `docs/codex/DIAGRAM_SOURCES.md` і можуть бути конвертовані в Mermaid, SVG або PNG під час фінального оформлення.": "У додатку перелічено діаграми, які вставлено в основні розділи роботи. Їх призначення полягає в узгодженому візуальному поданні архітектури АІНС, стану сесії та сценарію виконання прототипу.",
        "У поточному середовищі локальний `ollama` binary не був встановлений, тому прототип додатково отримав деградаційний режим. Це важливий практичний висновок: для дослідницької системи потрібно з самого початку передбачати резервну логіку, а не вважати зовнішнє середовище виконання гарантовано доступним [25].": "У поточному середовищі локальний `ollama` binary не був встановлений, тому прототип використовує чесно задокументований резервний режим. Після виправлень у коді цей режим доповнено явною schema validation, retry/fallback для JSON і дружнім повідомленням користувачу без ламання UX [25].",
        "Предмет дослідження - архітектурні моделі, методи адаптації великих мовних моделей, механізми пам'яті агента та програмні засоби реалізації vertical slice для навчання базового JavaScript.": "Предмет дослідження - архітектурні моделі, методи адаптації великих мовних моделей, механізми пам'яті агента та програмні засоби реалізації мінімального практичного сценарію навчання базового JavaScript.",
        "У роботі систематизовано класичні компоненти ITS, проаналізовано сучасні підходи до використання LLM в освіті, запропоновано архітектуру агента-репетитора з моделлю стану студента, матрицею рішень і багаторівневою організацією пам'яті, а також реалізовано мінімальний браузерний прототип із локальним LLM-середовище виконання, Web Worker sandbox та резервною логікою.": "У роботі систематизовано класичні компоненти ITS, проаналізовано сучасні підходи до використання LLM в освіті, запропоновано архітектуру агента-репетитора з моделлю стану студента, матрицею рішень і багаторівневою організацією пам'яті, а також реалізовано мінімальний браузерний прототип із локальною LLM, `Web Worker` для тестів та резервною логікою.",
        "Метою роботи є проєктування архітектури АІНС на основі великих мовних моделей для реалізації діалогового агента з функціями репетитора та демонстрація цієї архітектури на мінімальному, але реальному vertical slice прототипу.": "Метою роботи є проєктування архітектури АІНС на основі великих мовних моделей для реалізації діалогового агента з функціями репетитора та демонстрація цієї архітектури на мінімальному, але реальному практичному сценарії прототипу.",
        "реалізувати мінімальний browser-based прототип для курсового vertical slice з локальним LLM runtime та безпечним тестовим запуском коду;": "реалізувати мінімальний браузерний прототип для курсового практичного сценарію з локальною LLM та безпечним тестовим запуском коду;",
        "Ключове обмеження полягає в тому, що локальний Ollama середовище виконання не був встановлений у поточному середовищі. Це означає, що частина практичної логіки перевірялася у резервному режимі, а повна робота з реальною локальною моделлю потребує ручного запуску.": "Ключове обмеження полягає в тому, що локальне середовище виконання Ollama не було встановлене у поточному середовищі. Це означає, що частина практичної логіки перевірялася у резервному режимі, а повна робота з реальною локальною моделлю потребує ручного запуску.",
        "реалізувати мінімальний браузерний прототип для курсового vertical slice з локальним LLM runtime та безпечним тестовим запуском коду;": "реалізувати мінімальний браузерний прототип для курсового практичного сценарію з локальною LLM та безпечним тестовим запуском коду;",
        "У центрі персоналізації перебуває модель студента. Повноцінні ITS використовують для цього складніші probabilistic або knowledge tracing механізми, але для vertical slice доцільно застосувати спрощений вектор стану `UserState`, який усе ж зберігає ключові сигнали для педагогічного рішення [2], [3], [4].": "У центрі персоналізації перебуває модель студента. Повноцінні ITS використовують для цього складніші probabilistic або knowledge tracing механізми, але для мінімального практичного сценарію доцільно застосувати спрощений вектор стану `UserState`, який усе ж зберігає ключові сигнали для педагогічного рішення [2], [3], [4].",
        "Оскільки в репозиторії на початку роботи не було прикладної реалізації, практичну частину свідомо обмежено до одного мінімального vertical slice. Метою було не створити повну платформу, а отримати реальний кодовий артефакт, що демонструє архітектурні рішення з розділу 3.": "Оскільки в репозиторії на початку роботи не було прикладної реалізації, практичну частину свідомо обмежено до одного мінімального практичного сценарію. Метою було не створити повну платформу, а отримати реальний кодовий артефакт, що демонструє архітектурні рішення з розділу 3.",
        "Отже, інженерна функціональність vertical slice підтверджена частково: чисті модулі й базова статична доставка працюють, але повний інтерактивний сценарій з локальним Ollama має бути додатково перевірений вручну.": "Отже, інженерна функціональність мінімального практичного сценарію підтверджена частково: чисті модулі й базова статична доставка працюють, але повний інтерактивний сценарій з локальним Ollama має бути додатково перевірений вручну.",
        "На момент завершення автоматично згенерованої чернетки наявні лише результати інженерної валідації. Їх достатньо, щоб стверджувати: у репозиторії присутній реальний vertical slice, який структурно відповідає описаній архітектурі, але вони недостатні для висновків про педагогічну ефективність.": "На момент завершення автоматично згенерованої чернетки наявні лише результати інженерної валідації. Їх достатньо, щоб стверджувати: у репозиторії присутній реальний мінімальний практичний сценарій, який структурно відповідає описаній архітектурі, але вони недостатні для висновків про педагогічну ефективність.",
        "Ключове обмеження полягає в тому, що локальний Ollama runtime не був встановлений у поточному середовищі. Це означає, що частина практичної логіки перевірялася у резервному режимі, а повна робота з реальною локальною моделлю потребує ручного запуску.": "Ключове обмеження полягає в тому, що локальне середовище виконання Ollama не було встановлене у поточному середовищі. Це означає, що частина практичної логіки перевірялася у резервному режимі, а повна робота з реальною локальною моделлю потребує ручного запуску.",
        "Для практичної частини реалізовано мінімальний vertical slice JavaScript tutor: генерація вправи, введення коду, запуск тестів у Web Worker, адаптивний фідбек, SessionMemory та резервна логіка у випадку недоступності або нестабільності локальної LLM.": "Для практичної частини реалізовано мінімальний практичний сценарій JavaScript tutor: генерація вправи, введення коду, запуск тестів у `Web Worker`, адаптивний фідбек, `SessionMemory` та резервна логіка у випадку недоступності або нестабільності локальної LLM.",
        "3.5.3. Матриця рішень і minimal feedback strategy": "3.5.3. Матриця рішень і стратегія мінімальної підказки",
        "4.3. Розробка prompt-стратегій": "4.3. Проєктування промптів для генерації вправ і фідбеку",
    }

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text], align=paragraph.alignment)

    appendix_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "У додатку перелічено діаграми, які вставлено в основні розділи роботи. Їх призначення полягає в узгодженому візуальному поданні архітектури АІНС, стану сесії та сценарію виконання прототипу."
    )
    appendix_items = [
        "Рисунок 3.1 - Компонентна архітектура АІНС для сценарію навчання базового JavaScript.",
        "Рисунок 3.2 - Структура `UserState` та зв'язок із `SessionMemory`.",
        "Рисунок 3.3 - Послідовність дій у сценарії генерації вправи, тестування і фідбеку.",
        "Рисунок 4.1 - Схема виконання прототипу з локальною LLM, резервним режимом і `Web Worker`.",
    ]
    current = appendix_paragraph._p.getnext()
    while current is not None and not (
        current.tag.endswith("}p") and Paragraph(current, appendix_paragraph._parent).text.strip() == "Додаток Б. Приклади промптів"
    ):
        next_current = current.getnext()
        if current.tag.endswith("}p"):
            paragraph = Paragraph(current, appendix_paragraph._parent)
            if paragraph.text.strip():
                remove_paragraph(paragraph)
            else:
                appendix_paragraph._p.getparent().remove(current)
        else:
            appendix_paragraph._p.getparent().remove(current)
        current = next_current
    for item in reversed(appendix_items):
        paragraph = appendix_paragraph.insert_paragraph_after(item) if hasattr(appendix_paragraph, "insert_paragraph_after") else None
        if paragraph is None:
            paragraph = appendix_paragraph._parent.add_paragraph()
            appendix_paragraph._p.addnext(paragraph._p)
            set_paragraph_text(paragraph, item)


def ensure_appendix_list(document: Document) -> None:
    appendix_paragraph = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "Додаток А. Діаграми архітектури системи"
    )
    note = appendix_paragraph._p.getnext()
    if note is not None and note.tag.endswith("}p"):
        note_paragraph = Paragraph(note, appendix_paragraph._parent)
        set_paragraph_text(
            note_paragraph,
            "У додатку стисло перелічено рисунки, які вставлено в основні розділи роботи. Вони відображають архітектуру АІНС, структуру стану сесії, послідовність сценарію та схему виконання прототипу.",
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    items = [
        "Рисунок 3.1 - Компонентна архітектура АІНС для сценарію навчання базового JavaScript.",
        "Рисунок 3.2 - Структура `UserState` та зв'язок із `SessionMemory`.",
        "Рисунок 3.3 - Діаграма послідовності для сценарію «згенерувати вправу -> пройти тести -> дати фідбек».",
        "Рисунок 4.1 - Схема виконання прототипу з локальною LLM, резервним режимом і `Web Worker`.",
    ]

    current = note.getnext() if note is not None else None
    item_index = 0
    while current is not None and current.tag.endswith("}p"):
        paragraph = Paragraph(current, appendix_paragraph._parent)
        text = paragraph.text.strip()
        if text == "Додаток Б. Приклади промптів":
            break
        if item_index < len(items):
            set_paragraph_text(paragraph, items[item_index], align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            item_index += 1
        else:
            remove_paragraph(paragraph)
        current = current.getnext()


def update_appendix_code_blocks(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        replacement = None
        if 'if (repeatedError || attemptsCount >= 2)' in text:
            replacement = dedent(
                """\
                const currentSignature = buildRunErrorSignature(runResult);
                const previousSignature = latestHistoryEntry?.signature === currentSignature
                  ? userState.errorHistory?.[1]?.signature
                  : latestHistoryEntry?.signature;
                if (previousSignature === currentSignature || attemptsCount >= 3) {
                  return { action: "targeted_hint" };
                }
                """
            ).strip()
        elif 'const worker = new Worker(new URL("../workers/jsSandboxWorker.js", import.meta.url)' in text:
            replacement = dedent(
                """\
                const payload = { userCode, functionName, tests };
                const validation = validateSandboxRequest(payload);
                worker.addEventListener("error", handleWorkerError);
                worker.addEventListener("messageerror", handleWorkerMessageError);
                worker.postMessage(payload);
                """
            ).strip()
        elif "this.userState.errorHistory = [newError, ...this.userState.errorHistory].slice(0, 8);" in text:
            replacement = dedent(
                """\
                this.userState.attemptsCount += 1;
                this.userState.confidence = deriveConfidence(runResult);
                this.userState.errorHistory = [newError, ...this.userState.errorHistory].slice(0, 8);
                this.lastErrorSignature = buildRunErrorSignature(runResult);
                """
            ).strip()

        if replacement:
            set_paragraph_text(paragraph, replacement, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
            for run in paragraph.runs:
                set_run_font(run, size=10, font_name="Courier New")


def save_document(document: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def main() -> None:
    if not INPUT_PATH.exists():
        raise FileNotFoundError(f"Input document not found: {INPUT_PATH}")

    WORK_DIR.mkdir(parents=True, exist_ok=True)
    diagram_paths = generate_diagrams()

    document = Document(INPUT_PATH)
    normalize_document_styles(document)
    configure_page_numbers(document)
    update_textual_content(document)
    normalize_section_breaks(document)
    add_table_captions(document)
    update_core_tables(document)
    strengthen_experiment_section(document)
    attach_figures(document, diagram_paths)
    update_key_paragraphs(document)
    ensure_appendix_list(document)
    update_appendix_code_blocks(document)

    headings = collect_headings(document)
    replace_toc(document, headings, page_map=None)
    save_document(document, TEMP_STAGE_ONE)
    render_docx_to_pdf(TEMP_STAGE_ONE, TEMP_PDF_ONE)

    page_map = extract_page_map(TEMP_PDF_ONE, headings)
    replace_toc(document, headings, page_map=page_map)
    save_document(document, TEMP_STAGE_TWO)
    render_docx_to_pdf(TEMP_STAGE_TWO, TEMP_PDF_TWO)

    provisional_pages = len(PdfReader(str(TEMP_PDF_TWO)).pages)
    update_annotation_counts(document, provisional_pages)
    save_document(document, OUTPUT_PATH)
    render_docx_to_pdf(OUTPUT_PATH, FINAL_PDF)

    final_page_count = len(PdfReader(str(FINAL_PDF)).pages)
    if final_page_count != provisional_pages:
        update_annotation_counts(document, final_page_count)
        save_document(document, OUTPUT_PATH)
        render_docx_to_pdf(OUTPUT_PATH, FINAL_PDF)


if __name__ == "__main__":
    main()
