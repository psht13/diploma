# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import shutil
import subprocess
import time
from pathlib import Path

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOC = ROOT / "assets" / "coursework_draft_ua_submission_ready_template.docx"
PREFERRED_SOURCE = ROOT / "output" / "coursework_draft_ua_submission_ready_v5.docx"
FALLBACK_SOURCE = ROOT / "tmp" / "docs" / "submission_ready_v3" / "coursework_submission_stage_v3.docx"
TARGET_DOC = ROOT / "output" / "coursework_draft_ua_submission_ready_v6.docx"
RENDER_DIR = ROOT / "tmp" / "docs" / "submission_ready_v6_final"
PDF_PATH = RENDER_DIR / f"{TARGET_DOC.stem}.pdf"
PNG_DIR = RENDER_DIR / "pages"

TRACKED_STYLES = [
    "Normal",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "ListingCaption",
    "SourceNote",
    "PromptBlock",
    "CodeBlock",
    "DialogueBlock",
    "AppendixFigureLabel",
]

EARLY_FRONTMATTER_TITLES = {
    "АНОТАЦІЯ",
    "ABSTRACT",
    "ЗМІСТ",
}

PARAGRAPH_REPLACEMENTS = {
    "Реальні результати педагогічного експерименту у поточній версії відсутні; замість цього подано чесно оформлений дизайн експерименту та план оцінювання для подальшого дослідження.": "У роботі не проводилося емпіричне педагогічне оцінювання; натомість подано проєкт дизайну експерименту та план подальшого дослідження.",
    "The document separates conceptual architecture from practical implementation. A local Ollama-based execution environment, Web Worker sandboxing, JSON retry and backup logic, and an honest experimental design without fabricated educational outcomes are treated as core principles.": "The document separates conceptual architecture from practical implementation. Core elements include a local Ollama-based execution environment, Web Worker sandboxing, JSON retry and fallback logic, and a proposed design for subsequent pedagogical evaluation.",
    "Сучасний ландшафт ШІ-репетитор systems містить як комерційні продукти, так і дослідницькі системи. Продуктові системи демонструють життєздатність формату, але зазвичай мало розкривають внутрішню архітектуру; академічні праці, навпаки, частіше дають детальний опис модулів і методів оцінювання [6], [7], [8], [19], [21], [23], [24].": "Сучасний ландшафт ШІ-репетиторів охоплює як комерційні продукти, так і дослідницькі системи. Продуктові рішення демонструють життєздатність формату, але зазвичай мало розкривають внутрішню архітектуру; академічні праці, навпаки, частіше дають детальний опис модулів і методів оцінювання [6], [7], [8], [19], [21], [23], [24].",
    "4.4.1. Функціональна валідація прототипу": "4.4.1. Функціональна перевірка прототипу",
    "У прототипі ця модель навмисно спрощена до JSON-репрезентації вправи з полями title, prompt, starterCode, functionName, tests, concepts, rubric. Таке рішення дозволяє зберегти місце для майбутнього RAG, але не розширювати практичну частину понад необхідний мінімум.": "У прототипі ця модель навмисно спрощена до JSON-подання вправи, у якому фіксуються назва вправи, формулювання задачі, початковий код, назва функції, тестові приклади, перелік понять і критерії оцінювання. Таке рішення дозволяє зберегти місце для майбутнього RAG, але не розширювати практичну частину понад необхідний мінімум.",
    "S = K × E × C × Topic × Attempts × M": "S = K × E × C × Th × N × M",
    "R: S × A × Ctx → Text": "R: S × A × Cx → V",
    "Тут: K - інтегрована оцінка засвоєння теми; E - історія типових помилок; C - робоча оцінка впевненості; Topic - поточна тема; Attempts - кількість спроб; M - пам'ять сесії; O - нові спостереження; A - множина педагогічних дій; T - функція переходу стану; P - політика вибору дії; R - функція формування відповіді; Ctx - контекст генерації; Text - текстова відповідь агента. Такий запис відокремлює стан, політику та мовну генерацію і робить архітектуру агента придатною для формального опису без ототожнення її з однією лише LLM [11], [12].": "У цьому записі змінна стану охоплює інтегровану оцінку засвоєння теми (K), історію типових помилок (E), робочу оцінку впевненості (C), поточну тему навчання (Th), кількість спроб (N) і пам'ять сесії (M); O позначає нові спостереження, A - множину педагогічних дій, T - функцію переходу стану, P - політику вибору дії, R - функцію формування відповіді, а Cx і V - відповідно контекст відповіді та сформовану текстову відповідь агента. Такий запис відокремлює стан, політику та мовну генерацію і робить архітектуру агента придатною для формального опису без ототожнення її з однією лише LLM [11], [12].",
    "Такий обсяг реалізації відповідає принципу пріоритету архітектури: замість великої кількості другорядних функцій свідомо реалізовано лише ті модулі, які демонструють зв'язок між генерацією задачі, моделлю стану, тестовим запуском, режимом пояснення та адаптивним репетиторським фідбеком. На рівні TutorPolicy і main.js додано явний режим пояснення, тому натискання кнопки Пояснити тему завжди переводить агента до concept_explanation, а на рівні SessionMemory та UI реалізовано полегшений JSON-експорт поточної сесії як допоміжний інструмент для подальших досліджень типу A/B або кросовер-досліджень.": "Такий обсяг реалізації відповідає принципу пріоритету архітектури: замість великої кількості другорядних функцій свідомо реалізовано лише ті модулі, які демонструють зв'язок між генерацією задачі, моделлю стану, тестовим запуском, режимом пояснення та адаптивним репетиторським фідбеком. На рівні педагогічної політики та основного клієнтського сценарію додано явний режим пояснення, тому натискання кнопки «Пояснити тему» завжди переводить агента до концептуального пояснення, а на рівні пам'яті сесії та інтерфейсу користувача реалізовано експорт даних поточної сесії у форматі JSON як допоміжний інструмент для подальших досліджень типу A/B або кросовер-досліджень.",
    "У поточній роботі виконано лише ту функціональну перевірку, яку можна підтвердити реальними артефактами середовища: модульні тести (npm test), контрольну перевірку (npm run check), браузерний прогін через Playwright і реальний локальний виклик Ollama + llama3.1:8b. Освітні результати, UX-оцінки та ефективність навчання тут не вигадуються. Натомість підтверджено коректну ескалацію підказок, явний режим пояснення, генерацію вправи локальною моделлю, полегшений JSON-експорт сесії та стійкість прототипу до невалідних LLM-відповідей.": "У роботі виконано функціональну перевірку прототипу, що охоплювала модульні тести, сценарну функціональну перевірку, браузерну функціональну перевірку та перевірку роботи локальної моделі Ollama на базі Llama 3.1:8b. За її результатами підтверджено коректну ескалацію підказок, явний режим пояснення, генерацію вправи за допомогою локальної моделі, експорт даних сесії у форматі JSON та стійкість прототипу до невалідних відповідей LLM. Освітні результати, UX-оцінки та ефективність навчання в межах курсової роботи не оцінювалися.",
    "Отже, інженерна функціональність мінімального практичного сценарію підтверджена як модульними тестами, так і повним браузерним сценарієм із локальною Ollama. Під час цього прогону було виявлено й усунуто два локальні недоліки: збільшено тайм-аут LLM-виклику для llama3.1:8b і виправлено скидання firstFailure під час переходу до нової вправи.": "Отже, працездатність мінімального практичного сценарію підтверджено як модульними тестами, так і повним браузерним сценарієм з використанням локальної моделі Ollama. У процесі технічної перевірки було виявлено й усунуто два локальні недоліки: збільшено тайм-аут звернення до моделі llama3.1:8b і виправлено скидання ознаки першої зафіксованої помилки під час переходу до нової вправи.",
    "На момент завершення роботи наявні лише результати інженерної валідації: модульні тести, контрольний сценарій npm run check і браузерний прогін через Playwright з локальною Ollama + llama3.1:8b.": "На момент завершення роботи наявні лише результати технічної перевірки: модульні тести, сценарна функціональна перевірка, браузерна функціональна перевірка та перевірка роботи локальної моделі Ollama на базі Llama 3.1:8b.",
    "Документ побудовано чесно щодо емпіричної частини: реальні освітні результати не вигадувалися, а замість них подано інженерну функціональну валідацію та дизайн майбутнього педагогічного експерименту.": "У документі подано технічну перевірку працездатності прототипу та дизайн подальшого педагогічного експерименту; емпіричні освітні результати в межах цієї роботи не оцінювалися.",
    "Нижче подано скорочені, але автентичні фрагменти промптів, що реально використовуються у прототипі для генерації вправ і формування адаптивного фідбеку.": "Нижче подано скорочені фрагменти промптів, використаних у прототипі для генерації вправ і формування адаптивного фідбеку.",
    "У прототипі ця декомпозиція реалізована не повністю: роль модуля сприйняття виконують явний режим пояснення, представлений параметром interactionMode зі значенням explain, та прості текстові евристики для вільного запиту, а роль модуля генерації виконує локальна LLM або резервний механізм. Саме така схема використовується як основа повнішої системи в магістерській проєкції.": "У прототипі ця декомпозиція реалізована не повністю: роль модуля сприйняття виконують явний режим пояснення, представлений параметром interactionMode зі значенням explain, та прості текстові евристики для вільного запиту, а роль модуля генерації виконує локальна LLM або резервний механізм. Саме така схема використовується як основа повнішої системи в магістерському проєкті.",
    "[21] LearnLM Team, Eedi, Wang A. et al. ШІ-репетиторing can safely and effectively support students: An exploratory RCT in UK classrooms. arXiv preprint arXiv:2512.23633. 2025. URL: https://arxiv.org/abs/2512.23633.": "[21] LearnLM Team, Eedi, Wang A. et al. AI tutoring can safely and effectively support students: An exploratory RCT in UK classrooms. arXiv preprint arXiv:2512.23633. 2025. URL: https://arxiv.org/abs/2512.23633.",
}

CELL_REPLACEMENTS = {
    "Не вигадувати результати експериментів": "Не заявляти непідтверджених результатів експерименту",
    "Зберігає UserState, поточну вправу, історію повідомлень, першу зафіксовану помилку та формує полегшений JSON-експорт сесії.": "Зберігає UserState, поточну вправу, історію повідомлень, першу зафіксовану помилку та формує експорт даних сесії у форматі JSON.",
    "модульні тести (npm test)": "модульні тести",
    "модульні тести, контрольна перевірка (npm run check), Playwright": "модульні тести, сценарна функціональна перевірка, браузерна функціональна перевірка",
    "контрольна перевірка (npm run check)": "сценарна функціональна перевірка",
    "Playwright + локальна Ollama": "браузерна функціональна перевірка з використанням локальної моделі Ollama",
    "Перевірено null, масив, рядок і частковий об'єкт для JSON-пакетів вправи та фідбеку.": "Перевірено випадки null, масиву, рядка і часткового об'єкта для JSON-пакетів вправи та фідбеку.",
    "Перевірено missing function, timeout path, functionName validation і обробку помилок worker.": "Перевірено відсутність цільової функції, сценарій тайм-ауту, валідацію functionName і обробку помилок worker.",
    "Підтверджено поля topic, difficulty, attemptsCount, firstFailure, lastAction, transcript, finalRunStatus, timestamp; firstFailure тепер належить поточній вправі.": "Підтверджено поля topic, difficulty, attemptsCount, firstFailure, lastAction, transcript, finalRunStatus, timestamp; поле firstFailure тепер належить поточній вправі.",
    "Скрипт без важкого стеку швидко перевіряє потік пояснення та формування JSON-експорту сесії.": "Сценарна функціональна перевірка швидко охоплює потік пояснення та формування JSON-експорту сесії.",
    "Підтверджено: генерація вправи з джерелом ollama, дві послідовні невдалі спроби з ескалацією підказки, явний режим пояснення та завантаження JSON-експорту сесії.": "Підтверджено генерацію вправи з джерелом ollama, дві послідовні невдалі спроби з ескалацією підказки, явний режим пояснення та завантаження JSON-експорту сесії.",
}


def source_doc() -> Path:
    return PREFERRED_SOURCE if PREFERRED_SOURCE.exists() else FALLBACK_SOURCE


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("pPr"):
            continue
        paragraph._p.remove(child)


def set_run_font(
    run,
    *,
    font_name: str,
    size_pt: float,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    font_name: str | None = None,
    size_pt: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    if font_name and size_pt is not None:
        set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic)
    if align is not None:
        paragraph.alignment = align


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def iter_body_blocks(doc: Document):
    for child in doc.element.body:
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def first_heading_index(doc: Document) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "Heading 1":
            return index
    return len(doc.paragraphs)


def copy_style_from_template(template_doc: Document, target_doc: Document, style_name: str) -> None:
    if style_name not in template_doc.styles or style_name not in target_doc.styles:
        return
    template_style = template_doc.styles[style_name]
    target_style = target_doc.styles[style_name]
    target_style.font.name = template_style.font.name
    target_style.font.size = template_style.font.size
    target_style.font.bold = template_style.font.bold
    target_style.font.italic = template_style.font.italic
    target_style.paragraph_format.alignment = template_style.paragraph_format.alignment
    target_style.paragraph_format.first_line_indent = template_style.paragraph_format.first_line_indent
    target_style.paragraph_format.left_indent = template_style.paragraph_format.left_indent
    target_style.paragraph_format.right_indent = template_style.paragraph_format.right_indent
    target_style.paragraph_format.space_before = template_style.paragraph_format.space_before
    target_style.paragraph_format.space_after = template_style.paragraph_format.space_after
    target_style.paragraph_format.line_spacing = template_style.paragraph_format.line_spacing
    target_style.paragraph_format.keep_with_next = template_style.paragraph_format.keep_with_next
    target_style.paragraph_format.keep_together = template_style.paragraph_format.keep_together
    target_style.paragraph_format.page_break_before = template_style.paragraph_format.page_break_before
    if target_style.element.rPr is not None and template_style.element.rPr is not None:
        r_fonts = target_style.element.rPr.rFonts
        template_r_fonts = template_style.element.rPr.rFonts
        if r_fonts is not None and template_r_fonts is not None:
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                value = template_r_fonts.get(qn(f"w:{attr}"))
                if value:
                    r_fonts.set(qn(f"w:{attr}"), value)


def apply_template_baseline(template_doc: Document, doc: Document) -> None:
    template_section = template_doc.sections[0]
    section = doc.sections[0]
    section.page_width = template_section.page_width
    section.page_height = template_section.page_height
    section.left_margin = template_section.left_margin
    section.right_margin = template_section.right_margin
    section.top_margin = template_section.top_margin
    section.bottom_margin = template_section.bottom_margin
    section.header_distance = template_section.header_distance
    section.footer_distance = template_section.footer_distance
    section.different_first_page_header_footer = template_section.different_first_page_header_footer

    for style_name in TRACKED_STYLES:
        copy_style_from_template(template_doc, doc, style_name)


def replace_exact_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        new_text = PARAGRAPH_REPLACEMENTS.get(paragraph.text)
        if new_text is not None:
            set_paragraph_text(paragraph, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = CELL_REPLACEMENTS.get(paragraph.text)
                    if new_text is not None:
                        set_paragraph_text(paragraph, new_text)


def remove_student_state_abbreviation(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Модель стану студента - спрощене подання стану студента в прототипі":
            remove_paragraph(paragraph)
            return


def ensure_title_metrics_line(doc: Document) -> None:
    first_heading = first_heading_index(doc)
    specialty_paragraph = None
    for paragraph in doc.paragraphs[:first_heading]:
        if paragraph.text.startswith("Спеціальність"):
            specialty_paragraph = paragraph
            break
    if specialty_paragraph is None:
        return

    for paragraph in doc.paragraphs[:first_heading]:
        if paragraph.text.startswith("Магістерська курсова робота:"):
            return

    metrics_paragraph = specialty_paragraph.insert_paragraph_before("Магістерська курсова робота: 00 с., 4 рис., 13 табл., 4 дод., 30 джерел.")
    metrics_paragraph.style = "Normal"
    metrics_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metrics_paragraph.paragraph_format.first_line_indent = Mm(0)
    for run in metrics_paragraph.runs:
        set_run_font(run, font_name="Times New Roman", size_pt=14.0)


def normalize_normal_like_paragraph(paragraph: Paragraph, *, font_size: float = 14.0) -> None:
    for run in paragraph.runs:
        if not run.text:
            continue
        set_run_font(
            run,
            font_name="Times New Roman",
            size_pt=font_size,
            bold=bool(run.bold),
            italic=bool(run.italic),
        )


def normalize_special_paragraph(paragraph: Paragraph, *, font_name: str, size_pt: float, bold: bool | None = None, italic: bool | None = None) -> None:
    for run in paragraph.runs:
        if not run.text:
            continue
        set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "1")
        tr_pr.append(cant_split)


def set_row_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def normalize_document(doc: Document) -> None:
    first_heading = first_heading_index(doc)

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name
        is_title_page = index < first_heading
        is_toc_line = "\t" in paragraph.text
        is_table_caption = text.startswith("Таблиця ")
        is_figure_caption = text.startswith("Рисунок ")

        if is_title_page:
            if paragraph.text.startswith("Магістерська курсова робота:"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Mm(0)
                normalize_normal_like_paragraph(paragraph, font_size=14.0)
            continue

        if style_name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            normalize_special_paragraph(paragraph, font_name="Times New Roman", size_pt=14.0, bold=True)
            continue

        if style_name in {"Heading 2", "Heading 3"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            normalize_special_paragraph(paragraph, font_name="Times New Roman", size_pt=14.0, bold=True)
            continue

        if style_name == "ListingCaption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            normalize_special_paragraph(paragraph, font_name="Times New Roman", size_pt=12.0, bold=True, italic=False)
            continue

        if style_name == "SourceNote":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            normalize_special_paragraph(paragraph, font_name="Times New Roman", size_pt=11.0, bold=False, italic=True)
            continue

        if style_name == "CodeBlock":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.left_indent = Mm(0)
            normalize_special_paragraph(paragraph, font_name="Courier New", size_pt=9.5, bold=False, italic=False)
            continue

        if style_name == "PromptBlock":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.left_indent = Mm(0)
            normalize_special_paragraph(paragraph, font_name="Courier New", size_pt=10.5, bold=False, italic=False)
            continue

        if style_name == "DialogueBlock":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            normalize_special_paragraph(paragraph, font_name="Times New Roman", size_pt=14.0, bold=False, italic=False)
            continue

        if is_toc_line:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            normalize_normal_like_paragraph(paragraph, font_size=14.0)
            continue

        if is_table_caption:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Mm(0)
            normalize_normal_like_paragraph(paragraph, font_size=14.0)
            continue

        if is_figure_caption:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            normalize_normal_like_paragraph(paragraph, font_size=14.0)
            continue

        if style_name == "Normal" and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Mm(12.5)
            normalize_normal_like_paragraph(paragraph, font_size=14.0)
            continue

        if style_name == "Normal":
            normalize_normal_like_paragraph(paragraph, font_size=14.0)

    for table_index, table in enumerate(doc.tables):
        if table_index == 0:
            continue
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_index, row in enumerate(table.rows):
            set_row_cant_split(row)
            if row_index == 0:
                set_row_header(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.style = "Normal"
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.left_indent = Mm(0)
                    paragraph.paragraph_format.first_line_indent = Mm(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        if not run.text:
                            continue
                        set_run_font(
                            run,
                            font_name="Times New Roman",
                            size_pt=12.0,
                            bold=row_index == 0,
                            italic=bool(run.italic) if row_index != 0 else False,
                        )


def render_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    profile_dir = Path("/tmp") / f"lo_profile_codex_v6_{os.getpid()}_{int(time.time() * 1000)}"
    profile_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if generated != pdf_path:
        shutil.copy2(generated, pdf_path)


def export_pngs(pdf_path: Path, png_dir: Path) -> None:
    if png_dir.exists():
        shutil.rmtree(png_dir)
    png_dir.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(pdf_path)
    for index, page in enumerate(pdf, start=1):
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        pix.save(str(png_dir / f"page_{index:02d}.png"))
    pdf.close()


def normalize_space(text: str) -> str:
    return " ".join(text.split())


def count_metrics(doc: Document) -> dict[str, int]:
    figure_count = sum(1 for paragraph in doc.paragraphs if normalize_space(paragraph.text).startswith("Рисунок "))
    table_count = sum(1 for paragraph in doc.paragraphs if normalize_space(paragraph.text).startswith("Таблиця "))
    appendix_count = sum(
        1
        for paragraph in doc.paragraphs
        if paragraph.style.name.startswith("Heading") and normalize_space(paragraph.text).startswith("Додаток ")
    )
    reference_count = sum(1 for paragraph in doc.paragraphs if normalize_space(paragraph.text).startswith("["))
    return {
        "figures": figure_count,
        "tables": table_count,
        "appendices": appendix_count,
        "references": reference_count,
    }


def set_first_run_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def extract_page_texts(pdf_path: Path) -> tuple[list[str], int]:
    pdf = fitz.open(pdf_path)
    texts = [normalize_space(page.get_text("text")) for page in pdf]
    page_count = pdf.page_count
    pdf.close()
    return texts, page_count


def locate_title_page(title: str, page_texts: list[str]) -> int:
    matches = [index + 1 for index, page_text in enumerate(page_texts) if title in page_text]
    if not matches:
        raise ValueError(f"Title not found in rendered PDF: {title}")
    if title in EARLY_FRONTMATTER_TITLES:
        return min(matches)
    return max(matches)


def toc_titles(doc: Document) -> list[str]:
    titles = []
    in_toc = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "ЗМІСТ":
            in_toc = True
            continue
        if in_toc and paragraph.style.name == "Heading 1":
            break
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name not in {"Heading 1", "Heading 2", "Heading 3"}:
            continue
        if text == "ЗМІСТ":
            continue
        titles.append(text)
    return titles


def update_toc(doc: Document, page_texts: list[str]) -> bool:
    changed = False
    toc_heading_index = next((i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "ЗМІСТ"), None)
    if toc_heading_index is None:
        return False
    toc_line_indices = []
    for index in range(toc_heading_index + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[index]
        if paragraph.style.name == "Heading 1":
            break
        toc_line_indices.append(index)

    titles = toc_titles(doc)
    if len(toc_line_indices) != len(titles):
        raise ValueError(f"TOC line count mismatch: {len(toc_line_indices)} vs {len(titles)}")

    for index, title in zip(toc_line_indices, titles):
        actual_page = locate_title_page(title, page_texts)
        new_text = f"{title}\t{actual_page}"
        if doc.paragraphs[index].text != new_text:
            set_first_run_text(doc.paragraphs[index], new_text)
            changed = True
    return changed


def sync_counts(doc_path: Path, pdf_path: Path) -> tuple[dict[str, int], bool]:
    page_texts, page_count = extract_page_texts(pdf_path)
    doc = Document(doc_path)
    metrics = count_metrics(doc)
    changed = False

    ua_counts = (
        f"Магістерська курсова робота: {page_count} с., {metrics['figures']} рис., "
        f"{metrics['tables']} табл., {metrics['appendices']} дод., {metrics['references']} джерел."
    )
    en_counts = (
        f"Master's coursework project: {page_count} pages, {metrics['figures']} figures, "
        f"{metrics['tables']} tables, {metrics['appendices']} appendices, and {metrics['references']} references."
    )

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Магістерська курсова робота:") and paragraph.text != ua_counts:
            set_first_run_text(paragraph, ua_counts)
            changed = True
        elif paragraph.text.startswith("Master's coursework project:") and paragraph.text != en_counts:
            set_first_run_text(paragraph, en_counts)
            changed = True

    if update_toc(doc, page_texts):
        changed = True

    if changed:
        doc.save(doc_path)

    metrics["pages"] = page_count
    return metrics, changed


def build_doc() -> tuple[Path, dict[str, int]]:
    source = source_doc()
    TARGET_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(source)
    template_doc = Document(TEMPLATE_DOC)
    apply_template_baseline(template_doc, doc)
    replace_exact_paragraphs(doc)
    remove_student_state_abbreviation(doc)
    ensure_title_metrics_line(doc)
    normalize_document(doc)
    doc.save(TARGET_DOC)

    metrics = {"pages": 0, "figures": 0, "tables": 0, "appendices": 0, "references": 0}
    for _ in range(3):
        render_pdf(TARGET_DOC, PDF_PATH)
        metrics, changed = sync_counts(TARGET_DOC, PDF_PATH)
        if not changed:
            break
    render_pdf(TARGET_DOC, PDF_PATH)
    export_pngs(PDF_PATH, PNG_DIR)
    metrics, _ = sync_counts(TARGET_DOC, PDF_PATH)
    if _:
        render_pdf(TARGET_DOC, PDF_PATH)
        export_pngs(PDF_PATH, PNG_DIR)
        metrics, _ = sync_counts(TARGET_DOC, PDF_PATH)
    return source, metrics


def main() -> None:
    source, metrics = build_doc()
    print(f"source={source}")
    print(f"target={TARGET_DOC}")
    print(
        "metrics="
        f"pages:{metrics['pages']},figures:{metrics['figures']},tables:{metrics['tables']},"
        f"appendices:{metrics['appendices']},references:{metrics['references']}"
    )


if __name__ == "__main__":
    main()
